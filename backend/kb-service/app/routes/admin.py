"""
KB Service — 管理后台路由

提供 KBD/SOP 文档审核和发布接口。
仅供管理员使用，需 INTERNAL_API_TOKEN 鉴权。

POST /api/admin/kbd/{id}/approve  — KBD 条目审核通过（生成 embedding + tsv）
POST /api/admin/sop/{id}/approve  — SOP 文档审核通过（解析决策树）
"""

from __future__ import annotations

import asyncio
import copy
import hashlib
import io
import json
import re
from datetime import UTC, datetime
from typing import TYPE_CHECKING, Annotated, Any

import jsonschema
from fastapi import APIRouter, File, Form, HTTPException, Request, UploadFile, status
from fastapi.responses import JSONResponse
from pydantic import BaseModel, Field, StrictInt
from shared.dynamic_resource.adapters import kbd_resource_payload, sop_resource_payload
from shared.dynamic_resource.loader import snapshot_revision_metadata
from shared.dynamic_resource.publisher import DynamicResourcePublisher
from shared.models.skill_definition import SkillDefinitionORM
from shared.models.tool_definition import ToolDefinitionORM
from shared.observability.logger import get_logger
from shared.observability.otel import get_current_trace_id
from shared.schemas.acquirer_args import normalize_qfk_system_args, validate_acquire_args
from shared.schemas.capability_descriptor import capability_descriptor_document, get_capability_descriptor
from shared.schemas.kbd_signal_safety import validate_kbd_read_only_signals_json
from shared.schemas.signal_output import sync_signal_requires
from shared.schemas.signal_schema import (
    certify_publishable_signals_json,
    normalize_optional_matcher_nulls,
    validate_draft_signals_json,
    validate_publishable_signals_json,
)
from shared.schemas.signal_schema import (
    humanize_signal_validation_error as humanize_schema_validation_error,
)
from shared.schemas.verification_contract import (
    expert_editor_issues,
    normalize_legacy_role_contract,
    reconcile_verification_contract,
)
from shared.utils.acquisition_strategy import parse_strategy
from sqlalchemy import select, text

from app.models.kbd_entry import KbdEntry, build_kbd_embedding_text, strip_markdown
from app.models.kbd_revision import KbdRevision
from app.models.sop_document import SopDocument
from app.schemas.sop_template import ValidationIssue
from app.services.kbd_mutation_guard import PublishedKbdMutationError, require_mutable_kbd
from app.services.kbd_review_metadata import build_review_metadata, normalize_change_annotations
from app.services.kbd_revision_service import (
    KBD_PAYLOAD_FIELDS,
    apply_kbd_revision_payload,
    build_kbd_revision_payload,
    diff_revision_payloads,
    ensure_kbd_revision,
    ensure_kbd_revision_payload,
    is_evaluation_candidate,
    resolve_proposal_baseline,
    revision_metadata,
    select_current_expert_pair,
    summarize_expert_signal_changes,
)
from app.services.sop_parser import extract_sop_variables, merge_variable_schema, parse_sop_markdown
from app.services.sop_tool_contract_validator import validate_sop_tool_contract
from app.utils.jieba_hci import segment


def _load_signals_json(raw: Any) -> dict:
    """把存储态 signals_json（dict 或 JSON 字符串）解析为 v2 文档 dict。

    仅做 JSON 解析，不做任何 v1 归一——运行时仅存在 v2 单一版本。
    """
    if isinstance(raw, str):
        try:
            document = json.loads(raw)
        except json.JSONDecodeError:
            return {}
    else:
        document = raw or {}
    # 旧 UI 把 qfk_system.container=host 当成 Terminal Bridge 执行域。新契约
    # 中 container 只表示 aCLI --container；host 的等价表达是省略该字段。读取
    # 即归一，使专家下一次保存自然完成无损迁移，且旧 published 文档仍可消费。
    if isinstance(document, dict):
        for signal in document.get("signals") or []:
            acquire = signal.get("acquire") if isinstance(signal, dict) else None
            if not isinstance(acquire, dict) or acquire.get("tool") != "qfk_system":
                continue
            args = acquire.get("args")
            if isinstance(args, dict) and args.get("container") == "host":
                args.pop("container", None)
    return document if isinstance(document, dict) else {}


def _signals_for_response(raw: Any) -> dict:
    """GET 响应直出 v2 文档（前端原生读 v2 对象化，RFC §7 演进，2026-07-22）：

    直接返回标准化 v2 文档 ``{schema_version, signals:[...]}``，前端
    ``KbdReviewView.vue`` 已原生基于 v2 结构渲染/编辑（无适配层、零信息损失）。
    运行时仅存在 v2 单一版本，无 v1 扁平 list 与 to_legacy_signal 反向桥接。"""
    document = _load_signals_json(raw)
    normalized, _ = normalize_legacy_role_contract(document)
    return normalized


def _normalize_qfk_system_command_args(document: dict[str, Any]) -> dict[str, Any]:
    """在保存工作稿前收敛 qfk_system 的唯一命令表达。

    完整命令只做可逆 argv 规范化；旧 resource_keyword 不具备可推导命令语义，
    特别不能把 VM ID 盲目追加给 lsof，因此由共享契约要求人工复核。
    """

    for signal in document.get("signals") or []:
        acquire = signal.get("acquire") if isinstance(signal, dict) else None
        if not isinstance(acquire, dict) or acquire.get("tool") != "qfk_system":
            continue
        acquire["args"] = normalize_qfk_system_args(acquire.get("args") or {})
    return document


def _strip_legacy_expert_provenance_flags(document: dict[str, Any]) -> dict[str, Any]:
    """兼容移除旧 Admin UI 写入的非契约 provenance 标记。

    ``expert_created``/``expert_restored`` 曾由页面用于标记本地交互来源，
    却不属于 Signal ``provenance``（来源事实）的 Schema。它们因此会让恢复
    拒绝候选或复制信号的整个工作稿返回 422，反而堵住专家修复路径。专家操作
    本身已在 KbdRevision 的 review metadata 中审计，所以保存边界可安全移除
    这两个仅前端使用过的历史字段；其他未知 provenance 字段仍由 Schema 拒绝。
    """

    for signal in document.get("signals") or []:
        if not isinstance(signal, dict):
            continue
        provenance = signal.get("provenance")
        if not isinstance(provenance, dict):
            continue
        provenance.pop("expert_created", None)
        provenance.pop("expert_restored", None)
    return document


def _prepare_expert_publish_signals(raw: Any) -> dict[str, Any]:
    """规范化历史角色冲突，并用当前工具契约为 Expert 发布内容盖章。"""

    normalized, _ = normalize_legacy_role_contract(_load_signals_json(raw))
    _strip_legacy_expert_provenance_flags(normalized)
    normalize_optional_matcher_nulls(normalized)
    canonical, _ = reconcile_verification_contract(normalized)
    validate_kbd_read_only_signals_json(canonical)
    return certify_publishable_signals_json(canonical)


def _validate_kbd_draft_signals_json(raw: Any) -> None:
    """KBD 工作稿门禁：先给出处置边界原因，再校验通用 v2 结构。"""

    validate_kbd_read_only_signals_json(raw)
    validate_draft_signals_json(raw)


def _validate_kbd_publishable_signals_json(raw: Any) -> None:
    """KBD 发布预检：先给出处置边界原因，再校验通用发布契约。"""

    validate_kbd_read_only_signals_json(raw)
    validate_publishable_signals_json(raw)


def _humanize_signal_validation_error(error: jsonschema.ValidationError, signals: list[Any]) -> dict[str, Any]:
    """把发布门禁错误转换成专家可以直接处理的语言。

    JSON Schema 路径是工程实现细节，不能作为专家审核任务。这里保留机器可检索
    code，但只返回页面可理解的说明和下一步动作。
    """

    message = error.message
    if "至少需要 1 条必要信号" in message:
        return {
            "level": "error",
            "code": "NO_MUST_SIGNAL",
            "location": "关键信号",
            "message": "发布前请至少保留一条“必要证据”。",
            "action": {"type": "edit_signal_role"},
        }
    if "缺少稳定 id" in message or "signal id 重复" in message:
        return {
            "level": "error",
            "code": "SIGNAL_ID_INVALID",
            "location": "关键信号",
            "message": "有一条关键信号的内部标识异常，请删除后重新新增该信号。",
        }
    if "输入变量没有上游产出" in message or "变量依赖存在环或不可达" in message:
        return {
            "level": "error",
            "code": "SIGNAL_VARIABLE_DEPENDENCY_INVALID",
            "location": "关键信号 / 输入变量",
            "message": message,
        }
    if "处置动作不属于 KBD 关键信号" in message:
        index_match = re.search(r"signals\[(\d+)]", message)
        signal_index = int(index_match.group(1)) if index_match else -1
        signal_id = (
            str(signals[signal_index].get("id") or "").strip()
            if 0 <= signal_index < len(signals) and isinstance(signals[signal_index], dict)
            else ""
        )
        issue: dict[str, Any] = {
            "level": "error",
            "code": "KBD_SOLUTION_SIGNAL_FORBIDDEN",
            "location": f"关键信号 · {signal_id} · 执行阶段" if signal_id else "关键信号 / 执行阶段",
            "message": "处置动作不属于 KBD 关键信号；请保留在解决方案中，或将该信号改为只读诊断检查。",
        }
        if signal_id:
            issue.update(
                {
                    "signal_id": signal_id,
                    "action": {"type": "edit_signal", "signal_id": signal_id, "focus": "orchestrate.phase"},
                }
            )
        return issue
    return humanize_schema_validation_error(error, signals)


def _raise_signal_validation_error(error: jsonschema.ValidationError, signals: list[Any]) -> None:
    """统一把保存、删除和发布错误返回为前端可定位的结构化问题。"""

    raise HTTPException(
        status_code=422,
        detail=_humanize_signal_validation_error(error, signals),
    ) from error


def _prepare_expert_draft_signals(raw: Any) -> dict[str, Any]:
    """归约并校验专家工作稿，保存与按 ID 删除共用同一权威边界。"""

    document = copy.deepcopy(_load_signals_json(raw))
    _strip_legacy_expert_provenance_flags(document)
    normalize_optional_matcher_nulls(document)
    document, _ = reconcile_verification_contract(document)
    try:
        document = _normalize_qfk_system_command_args(document)
    except ValueError as exc:
        raise HTTPException(
            status_code=422,
            detail={
                "level": "error",
                "code": "SIGNAL_ACQUIRE_ARGS_INVALID",
                "location": "关键信号 / 采集参数",
                "message": str(exc),
            },
        ) from exc
    metadata = document.get("generation_metadata")
    if isinstance(metadata, dict):
        metadata["status"] = "manual_reviewed"
    signals = document.get("signals") or []
    for signal in signals:
        if not isinstance(signal, dict):
            continue
        sync_signal_requires(signal)
        acquire = signal.get("acquire") or {}
        ok, error = validate_acquire_args(acquire.get("tool"), acquire.get("args", {}))
        if not ok:
            signal_id = str(signal.get("id") or "").strip()
            raise HTTPException(
                status_code=422,
                detail={
                    "level": "error",
                    "code": "SIGNAL_ACQUIRE_ARGS_INVALID",
                    "location": f"关键信号 · {signal_id} · 采集参数" if signal_id else "关键信号 / 采集参数",
                    "message": str(error),
                    **(
                        {
                            "signal_id": signal_id,
                            "action": {"type": "edit_signal", "signal_id": signal_id, "focus": "acquire.args"},
                        }
                        if signal_id
                        else {}
                    ),
                },
            )
    try:
        _validate_kbd_draft_signals_json(document)
    except jsonschema.ValidationError as exc:
        _raise_signal_validation_error(exc, signals)
    return document


def _delete_signal_from_document(raw: Any, signal_id: str) -> dict[str, Any]:
    """按稳定 ID 删除权威工作稿中的 Signal，并同步重算 Agent Contract。"""

    normalized_id = str(signal_id or "").strip()
    document = copy.deepcopy(_load_signals_json(raw))
    signals = document.get("signals") or []
    if not normalized_id or not any(
        isinstance(signal, dict) and str(signal.get("id") or "") == normalized_id for signal in signals
    ):
        raise HTTPException(
            status_code=404,
            detail={
                "code": "KBD_SIGNAL_NOT_FOUND",
                "message": "目标关键信号已不存在，请刷新后重试。",
                "signal_id": normalized_id or None,
            },
        )
    document["signals"] = [
        signal
        for signal in signals
        if not (isinstance(signal, dict) and str(signal.get("id") or "") == normalized_id)
    ]
    return _prepare_expert_draft_signals(document)


if TYPE_CHECKING:
    from shared.database.postgres import DatabaseManager

    from app.schemas.sop_template import SOPNode
    from app.services.embedding import EmbeddingService

logger = get_logger("kb-service-admin")
router = APIRouter(prefix="/api/kb", tags=["admin"])

# 新增 KBD 审核路由（独立 prefix）
kbd_router = APIRouter(prefix="/api/admin/kbd", tags=["kbd-admin"])

# 新增 SOP 审核路由（独立 prefix）
sop_router = APIRouter(prefix="/api/admin/sop", tags=["sop-admin"])

_db_manager: DatabaseManager | None = None
_embedding_service: EmbeddingService | None = None


def set_dependencies(db: DatabaseManager, embedding: EmbeddingService | None = None) -> None:
    """注入数据库和 embedding 服务依赖"""
    global _db_manager, _embedding_service
    _db_manager = db
    _embedding_service = embedding


@kbd_router.get("/capabilities")
async def get_kbd_capabilities(request: Request) -> dict[str, Any]:
    """返回由代码契约生成的只读 Capability Descriptor。

    当前只证明 shared 参数契约已声明；不会把 Admin 数据库配置冒充为 Agent Handler
    已部署。后续可在相同信封中合并 agent-service 的运行时探测结果。
    """

    _check_auth(request)
    return capability_descriptor_document()


def _derive_capability_gaps(raw_signals: Any) -> list[dict[str, str]]:
    """从当前 Signal 契约派生工程能力缺口，不把部署状态伪装成专家待办。"""

    document = _load_signals_json(raw_signals)
    gaps: list[dict[str, str]] = []
    for signal in document.get("signals") or []:
        if not isinstance(signal, dict):
            continue
        tool = str((signal.get("acquire") or {}).get("tool") or "").strip()
        if not tool:
            continue
        descriptor = get_capability_descriptor(tool)
        if descriptor is None:
            gaps.append(
                {
                    "capability_id": tool,
                    "code": "CAPABILITY_UNDECLARED",
                    "signal_id": str(signal.get("id") or ""),
                    "message": "当前平台没有声明该采集能力；应实现能力或修改信号，不能靠自由 Shell 兜底。",
                }
            )
        elif descriptor["runtime_status"] != "available":
            gaps.append(
                {
                    "capability_id": tool,
                    "code": "CAPABILITY_RUNTIME_UNVERIFIED",
                    "signal_id": str(signal.get("id") or ""),
                    "message": "参数契约已声明，但 Agent Handler/部署状态尚未被运行时探测确认。",
                }
            )
    return gaps


@kbd_router.get("/capability-gaps")
async def get_kbd_capability_gaps(request: Request, limit: int = 200) -> dict[str, Any]:
    """提供工程侧 Capability Gap 聚合；不展示在专家的常规编辑待办中。"""

    _check_auth(request)
    if _db_manager is None:
        raise HTTPException(status_code=503, detail="数据库未就绪")
    bounded_limit = min(max(limit, 1), 1000)
    async with _db_manager.async_session_factory() as session:
        result = await session.execute(
            select(KbdEntry.id, KbdEntry.support_id, KbdEntry.title, KbdEntry.status, KbdEntry.signals_json)
            .order_by(KbdEntry.updated_at.desc())
            .limit(bounded_limit)
        )
        rows = result.all()
    grouped: dict[tuple[str, str], dict[str, Any]] = {}
    for row in rows:
        for gap in _derive_capability_gaps(row.signals_json):
            key = (gap["capability_id"], gap["code"])
            record = grouped.setdefault(
                key,
                {
                    "capability_id": gap["capability_id"],
                    "code": gap["code"],
                    "affected_kbd_count": 0,
                    "affected_signal_count": 0,
                    "examples": [],
                    "_kbd_ids": set(),
                },
            )
            record["affected_signal_count"] += 1
            if row.id not in record["_kbd_ids"]:
                record["_kbd_ids"].add(row.id)
                record["affected_kbd_count"] += 1
                if len(record["examples"]) < 5:
                    record["examples"].append(
                        {
                            "kbd_id": row.id,
                            "support_id": row.support_id,
                            "title": row.title,
                            "status": row.status,
                            "signal_id": gap["signal_id"],
                        }
                    )
    gaps = sorted(grouped.values(), key=lambda item: (-item["affected_signal_count"], item["capability_id"]))
    for item in gaps:
        item.pop("_kbd_ids", None)
    return {
        "schema_version": 1,
        "scanned_kbd_count": len(rows),
        "gaps": gaps,
    }


@kbd_router.get("/evaluation-export")
async def export_kbd_evaluation_data(request: Request, limit: int = 200) -> dict[str, Any]:
    """导出 LLM Proposal→Expert 与运行结果的结构化样本，不运行 Challenger。"""

    _check_auth(request)
    if _db_manager is None:
        raise HTTPException(status_code=503, detail="数据库未就绪")
    bounded_limit = min(max(limit, 1), 1000)
    async with _db_manager.async_session_factory() as session:
        revision_result = await session.execute(
            select(KbdRevision)
            .order_by(KbdRevision.kbd_entry_id, KbdRevision.revision_no)
            .limit(bounded_limit * 8)
        )
        revisions = list(revision_result.scalars().all())
        audit_result = await session.execute(
            text(
                """
                SELECT resource_name, revision, status, error, metadata_json, created_at
                FROM dynamic_resource_usage_audit
                WHERE resource_type = 'kbd'
                  AND consumer = 'agent-service.kbd_differential'
                ORDER BY created_at DESC
                LIMIT :limit
                """
            ),
            {"limit": bounded_limit * 20},
        )
        audits = list(audit_result.mappings().all())
    revisions_by_kbd: dict[int, list[KbdRevision]] = {}
    for revision in revisions:
        revisions_by_kbd.setdefault(int(revision.kbd_entry_id), []).append(revision)
    audits_by_resource: dict[tuple[str, int], list[dict[str, Any]]] = {}
    for audit in audits:
        audits_by_resource.setdefault((str(audit["resource_name"]), int(audit["revision"])), []).append(
            {
                "status": audit["status"],
                "error": audit["error"],
                "metadata": audit["metadata_json"] or {},
                "created_at": audit["created_at"].isoformat() if audit["created_at"] else None,
            }
        )
    examples: list[dict[str, Any]] = []
    for kbd_id, history in revisions_by_kbd.items():
        if len(examples) >= bounded_limit:
            break
        by_id = {item.id: item for item in history}
        for expert in (item for item in history if is_evaluation_candidate(item)):
            proposal = resolve_proposal_baseline(expert, by_id)
            if proposal is None:
                continue
            proposal_diff = diff_revision_payloads(proposal.payload_json, expert.payload_json)
            examples.append(
                {
                    "kbd_id": kbd_id,
                    "proposal_revision": {
                        "id": proposal.id,
                        "payload": proposal.payload_json,
                        "generation_metadata": proposal.generation_metadata or {},
                    },
                    "expert_revision": {
                        "id": expert.id,
                        "baseline_proposal_revision_id": proposal.id,
                        "payload": expert.payload_json,
                        "review_metadata": expert.review_metadata or {},
                        "validation_summary": expert.validation_summary or {},
                    },
                    "diff_from_proposal": proposal_diff,
                    # v1 消费方兼容别名；语义已固定为 Proposal baseline，不再是任意 direct parent。
                    "diff_from_parent": proposal_diff,
                    "runtime_outcomes": [
                        item
                        for (resource_name, _revision), items in audits_by_resource.items()
                        if resource_name == str(kbd_id)
                        for item in items
                    ],
                }
            )
            if len(examples) >= bounded_limit:
                break
    return {
        "schema_version": 2,
        "kind": "kbd_proposal_expert_runtime_export",
        "facts_boundary": {
            "trusted_reviewer": False,
            "expert_gold": False,
            "evidence_replay": False,
            "execution_replay": False,
            "real_customer_execution": False,
            "challenger_generated": False,
        },
        "examples": examples,
    }


@kbd_router.get("/runtime-metrics")
async def get_kbd_runtime_metrics(request: Request, limit: int = 5000) -> dict[str, Any]:
    """返回已发布 KBD 的实际消费、编译和 Signal 失败模式聚合。"""

    _check_auth(request)
    if _db_manager is None:
        raise HTTPException(status_code=503, detail="数据库未就绪")
    bounded_limit = min(max(limit, 1), 20000)
    async with _db_manager.async_session_factory() as session:
        result = await session.execute(
            text(
                """
                SELECT status, metadata_json
                FROM dynamic_resource_usage_audit
                WHERE resource_type = 'kbd'
                  AND consumer = 'agent-service.kbd_differential'
                ORDER BY created_at DESC
                LIMIT :limit
                """
            ),
            {"limit": bounded_limit},
        )
        rows = list(result.mappings().all())
    status_counts: dict[str, int] = {}
    candidate_state_counts: dict[str, int] = {}
    compile_counts: dict[str, int] = {}
    outcome_counts: dict[str, int] = {}
    failure_mode_counts: dict[str, int] = {}
    for row in rows:
        status_counts[row["status"]] = status_counts.get(row["status"], 0) + 1
        metadata = row["metadata_json"] or {}
        candidate_state = str(metadata.get("candidate_state") or "UNKNOWN")
        candidate_state_counts[candidate_state] = candidate_state_counts.get(candidate_state, 0) + 1
        compile_status = str((metadata.get("compile") or {}).get("status") or "not_recorded")
        compile_counts[compile_status] = compile_counts.get(compile_status, 0) + 1
        for item in metadata.get("signal_outcomes") or []:
            if not isinstance(item, dict):
                continue
            outcome = str(item.get("outcome") or "NOT_RECORDED")
            outcome_counts[outcome] = outcome_counts.get(outcome, 0) + 1
            failure_mode = item.get("failure_mode")
            if failure_mode:
                key = str(failure_mode)
                failure_mode_counts[key] = failure_mode_counts.get(key, 0) + 1
    return {
        "schema_version": 1,
        "sample_size": len(rows),
        "status_counts": status_counts,
        "candidate_state_counts": candidate_state_counts,
        "compile_counts": compile_counts,
        "signal_outcome_counts": outcome_counts,
        "failure_mode_counts": failure_mode_counts,
        "facts_boundary": "仅聚合 Agent 实际运行产生的审计；不等同于真实客户环境回放或 Expert Gold。",
    }


def _check_auth(request: Request) -> None:
    """验证内部服务 Token"""
    from app.config import settings

    auth_header = request.headers.get("Authorization", "")
    if not auth_header.startswith("Bearer "):
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="缺少 Bearer Token")
    token = auth_header.split(" ", 1)[1]
    if token != settings.INTERNAL_API_TOKEN:
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="Token 无效")


async def _require_directly_mutable_kbd(session, kbd_id: int, *, for_update: bool = False) -> KbdEntry:
    """把领域写入门禁转换为稳定的 Admin API 错误。"""

    try:
        return await require_mutable_kbd(session, kbd_id, for_update=for_update)
    except LookupError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    except PublishedKbdMutationError as exc:
        raise HTTPException(
            status_code=409,
            detail={
                "code": "KBD_MAINTENANCE_WORKING_REQUIRED",
                "message": str(exc),
                "agent_active_unchanged": True,
            },
        ) from exc


async def _publish_kbd_revision(session, kbd_id: int, trace_id: str | None) -> dict | None:
    """将 KBD 条目发布为动态资源 revision。"""
    from app.models.kbd_entry import KbdEntry

    result = await session.execute(select(KbdEntry).where(KbdEntry.id == kbd_id))
    kbd = result.scalar_one_or_none()
    if kbd is None:
        return None
    snapshot = await DynamicResourcePublisher(session).ensure_published(**kbd_resource_payload(kbd), trace_id=trace_id)
    return snapshot_revision_metadata(snapshot)


async def _freeze_approved_expert_revision(
    session,
    *,
    kbd_id: int,
    reviewer_id: int,
    review_note: str | None,
    trace_id: str | None,
) -> KbdRevision:
    """在发布事务内冻结通过门禁的 Expert Revision。"""

    kbd_result = await session.execute(select(KbdEntry).where(KbdEntry.id == kbd_id).with_for_update())
    kbd = kbd_result.scalar_one()
    if kbd.latest_proposal_revision_id is None:
        proposal_revision = await ensure_kbd_revision(
            session,
            kbd=kbd,
            revision_type="proposal",
            actor_type="migration",
            generation_metadata={"origin": "pre_review_baseline", "identity_status": "not_applicable"},
            trace_id=trace_id,
        )
    else:
        proposal_revision = await session.get(KbdRevision, kbd.latest_proposal_revision_id)
    parent_revision_id = kbd.working_revision_id or (
        proposal_revision.id if proposal_revision is not None else None
    )
    parent_revision = await session.get(KbdRevision, parent_revision_id) if parent_revision_id else None
    review_metadata = build_review_metadata(
        parent_payload=getattr(parent_revision, "payload_json", None) if parent_revision is not None else {},
        payload=build_kbd_revision_payload(kbd),
        identity_status="unverified_body_reviewer_id",
        review_state="approved",
        reviewer_id=reviewer_id,
        review_note=review_note,
    )
    return await ensure_kbd_revision(
        session,
        kbd=kbd,
        revision_type="expert",
        actor_type="expert",
        actor_id=None,
        parent_revision_id=parent_revision_id,
        generation_metadata={
            "origin": "admin_review",
            "identity_status": "unverified_body_reviewer_id",
            "legacy_reviewer_id": reviewer_id,
            "review_note": review_note or "",
        },
        validation_summary={"status": "passed", "gate": "publishable_signals_json"},
        review_metadata=review_metadata,
        trace_id=trace_id,
        # 即使知识 payload 未变化，也必须单独冻结“专家已批准”这一事实；工作稿 revision
        # 保持不可变，批准备注和 validation 不能被幂等复用吞掉。
        reuse_existing=False,
    )


async def _publish_sop_revision(session, document_id: int, trace_id: str | None) -> dict | None:
    """将 SOP 文档发布为动态资源 revision。"""
    result = await session.execute(select(SopDocument).where(SopDocument.id == document_id))
    sop = result.scalar_one_or_none()
    if sop is None:
        return None
    snapshot = await DynamicResourcePublisher(session).ensure_published(**sop_resource_payload(sop), trace_id=trace_id)
    return snapshot_revision_metadata(snapshot)


class DocumentUpdateRequest(BaseModel):
    """文档状态更新请求（已废弃，保留向后兼容）"""

    status: str | None = None  # draft/under_review/approved/published/rejected/archived
    review_note: str | None = None
    reviewer: str | None = None


# ─────────────────────────────────────────────────────────────────────────────
# KBD 条目列表查询接口（kbd_entry 表）
# ─────────────────────────────────────────────────────────────────────────────


@kbd_router.get("/pending/stats")
async def kbd_category_stats(
    request: Request,
    status: str = "",
):
    """查询各分类条目数量统计（供前端 Tab 分组展示使用）"""
    _check_auth(request)

    if _db_manager is None:
        raise HTTPException(status_code=503, detail="数据库未就绪")

    async with _db_manager.async_session_factory() as session:
        where_clause = "WHERE e.status = :status" if status else ""
        sql = text(f"""
            SELECT
                COALESCE(e.ai_category_id, '__uncategorized__') AS category_id,
                c.name AS category_name,
                COUNT(*) AS cnt
            FROM kbd_entry e
            LEFT JOIN kb_category c ON c.code = e.ai_category_id
            {where_clause}
            GROUP BY e.ai_category_id, c.name
            ORDER BY cnt DESC
        """)
        params_db: dict = {}
        if status:
            params_db["status"] = status
        result = await session.execute(sql, params_db)
        rows = result.mappings().all()

    return {
        "stats": [
            {
                "category_id": row["category_id"],
                "category_label": "未分类"
                if row["category_id"] == "__uncategorized__"
                else (f"{row['category_id']} {row['category_name']}" if row["category_name"] else row["category_id"]),
                "count": row["cnt"],
            }
            for row in rows
        ],
        "total": sum(r["cnt"] for r in rows),
    }


@kbd_router.get("/pending")
async def list_kbd_entries(
    request: Request,
    page: int = 1,
    page_size: int = 20,
    status: str = "",
    category_id: str | None = None,
    support_id: str | None = None,
    title_keyword: str | None = None,
    min_confidence: float | None = None,
    max_confidence: float | None = None,
    sort_by: str = "updated_at",
    sort_order: str = "desc",
):
    """查询 KBD 条目列表（分页 + 状态/分类/案例ID/标题/置信度过滤 + 排序）

    Args:
        page: 页码（从 1 开始）
        page_size: 每页条数（最大 100）
        status: 状态过滤（draft/published/rejected/archived）
        category_id: 按 AI 分类 ID 过滤（可选）
        support_id: 按案例 ID 精准匹配（可选）
        title_keyword: 按标题关键字模糊搜索（可选）
        min_confidence: 最低置信度（可选，0-1）
        max_confidence: 最高置信度（可选，0-1）
        sort_by: 排序字段（support_id/ai_category_conf/status/updated_at/created_at）
        sort_order: 排序方向（asc/desc）

    Returns:
        { entries: [...], total, page, page_size }
    """
    # 排序字段白名单（防 SQL 注入）
    valid_sort_columns = {"support_id", "ai_category_id", "ai_category_conf", "status", "updated_at", "created_at"}
    sort_column = sort_by if sort_by in valid_sort_columns else "updated_at"
    sort_dir = "DESC" if sort_order.lower() == "desc" else "ASC"
    _check_auth(request)

    if _db_manager is None:
        raise HTTPException(status_code=503, detail="数据库未就绪")

    # 参数校验
    page_size = min(max(page_size, 1), 100)
    page = max(page, 1)
    offset = (page - 1) * page_size

    logger.info(
        event="kbd_list_request",
        page=page,
        page_size=page_size,
        status=status,
        category_id=category_id,
        support_id=support_id,
        title_keyword=title_keyword,
    )

    async with _db_manager.async_session_factory() as session:
        # 构建 WHERE 条件
        where_clauses: list[str] = []
        params: dict = {"limit": page_size, "offset": offset}
        if status:
            where_clauses.append("status = :status")
            params["status"] = status

        if category_id:
            if category_id == "__uncategorized__":
                where_clauses.append("ai_category_id IS NULL")
            else:
                where_clauses.append("(ai_category_id = :category_id OR category_id = :category_id)")
                params["category_id"] = category_id

        # 按案例 ID 精准匹配
        if support_id:
            where_clauses.append("support_id = :support_id")
            params["support_id"] = support_id

        # 按标题关键字模糊搜索
        if title_keyword:
            where_clauses.append("title ILIKE :title_keyword")
            params["title_keyword"] = f"%{title_keyword}%"

        # 按置信度范围过滤
        if min_confidence is not None:
            where_clauses.append("ai_category_conf >= :min_confidence")
            params["min_confidence"] = min_confidence
        if max_confidence is not None:
            where_clauses.append("ai_category_conf <= :max_confidence")
            params["max_confidence"] = max_confidence

        where_sql = ("WHERE " + " AND ".join(where_clauses)) if where_clauses else ""

        # 查询总数
        count_sql = text(f"SELECT COUNT(*) FROM kbd_entry {where_sql}")  # noqa: S608
        count_result = await session.execute(count_sql, params)
        total = count_result.scalar() or 0

        # 查询分页数据
        data_sql = text(  # noqa: S608
            f"""
            SELECT e.id, e.support_id, e.title,
                   e.problem_description, e.alert_info, e.steps_text, e.root_cause,
                   e.solution, e.operational_impact, e.is_temporary, e.recommendations,
                   e.signals_json, e.content_md, e.content_raw, e.images_json,
                   e.metadata, e.category_id, e.ai_category_id,
                   e.ai_category_conf, e.ai_category_reason,
                   e.status, e.reviewer_id, e.review_note,
                   e.hit_count, e.created_at, e.updated_at,
                   c.name AS ai_category_name
            FROM kbd_entry e
            LEFT JOIN kb_category c ON c.code = e.ai_category_id
            {where_sql}
            ORDER BY {sort_column} {sort_dir}, e.id DESC
            LIMIT :limit OFFSET :offset
            """
        )
        result = await session.execute(data_sql, params)
        rows = result.mappings().all()

    entries = [
        {
            "id": row["id"],
            "support_id": row["support_id"],
            "title": row["title"],
            "problem_description": row["problem_description"] or "",
            "alert_info": row["alert_info"] or "",
            "steps_text": row["steps_text"] or "",
            "root_cause": row["root_cause"] or "",
            "solution": row["solution"] or "",
            "operational_impact": row["operational_impact"] or "",
            "is_temporary": row["is_temporary"] or "",
            "recommendations": row["recommendations"] or "",
            "signals_json": _signals_for_response(row["signals_json"]),
            "content_md": row["content_md"] or "",
            "content_raw": row["content_raw"] or "",
            "images_json": row["images_json"] or [],
            "metadata": row["metadata"] or {},
            "category_id": row["category_id"],
            "ai_category_id": row["ai_category_id"],
            "ai_category_label": f"{row['ai_category_id']} {row['ai_category_name']}"
            if row.get("ai_category_name") and row.get("ai_category_id")
            else (row["ai_category_id"] or ""),
            "ai_category_conf": float(row["ai_category_conf"]) if row["ai_category_conf"] is not None else None,
            "ai_category_reason": row["ai_category_reason"],
            "status": row["status"],
            "reviewer_id": row["reviewer_id"],
            "review_note": row["review_note"],
            "hit_count": row.get("hit_count", 0),
            "created_at": row["created_at"].isoformat() if row["created_at"] else None,
            "updated_at": row["updated_at"].isoformat() if row["updated_at"] else None,
        }
        for row in rows
    ]

    logger.info(event="kbd_list_response", total=total, returned=len(entries))

    return {
        "entries": entries,
        "total": total,
        "page": page,
        "page_size": page_size,
    }


# ─────────────────────────────────────────────────────────────────────────────
# KBD 条目单条详情接口
# ─────────────────────────────────────────────────────────────────────────────


@kbd_router.get("/{kbd_id}")
async def get_kbd_entry_detail(request: Request, kbd_id: int):
    """获取单个 KBD 条目详情（含完整 content_md）

    Args:
        kbd_id: KBD 条目 ID

    Returns:
        KBD 条目完整详情（含 content_md、metadata 等）
    """
    _check_auth(request)

    if _db_manager is None:
        raise HTTPException(status_code=503, detail="数据库未就绪")

    logger.info(event="kbd_detail_request", kbd_id=kbd_id, trace_id=get_current_trace_id())

    async with _db_manager.async_session_factory() as session:
        result = await session.execute(
            text("""
                SELECT id, support_id, title,
                       problem_description, alert_info, steps_text, root_cause,
                       solution, operational_impact, is_temporary, recommendations,
                       signals_json, content_md, content_raw, images_json,
                       metadata, category_id, ai_category_id,
                       ai_category_conf, ai_category_reason,
                       status, reviewer_id, review_note,
                       latest_proposal_revision_id, working_revision_id, lock_version,
                       created_at, updated_at, published_at
                FROM kbd_entry
                WHERE id = :id
            """),
            {"id": kbd_id},
        )
        row = result.mappings().first()

        if not row:
            raise HTTPException(status_code=404, detail=f"KBD 条目 {kbd_id} 不存在")

        working_payload: dict[str, Any] | None = None
        if row["status"] == "published" and row["working_revision_id"] is not None:
            working_revision = await session.get(KbdRevision, row["working_revision_id"])
            if (
                working_revision is not None
                and (working_revision.generation_metadata or {}).get("origin") == "admin_maintenance"
            ):
                working_payload = working_revision.payload_json

    def review_value(field: str, fallback: Any = None) -> Any:
        if working_payload is not None and field in working_payload:
            return working_payload[field]
        return row.get(field, fallback)

    return {
        "id": row["id"],
        "support_id": row["support_id"],
        "title": review_value("title"),
        # 8 大章节字段（完整内容）
        "problem_description": review_value("problem_description") or "",
        "alert_info": review_value("alert_info") or "",
        "steps_text": review_value("steps_text") or "",
        "root_cause": review_value("root_cause") or "",
        "solution": review_value("solution") or "",
        "operational_impact": review_value("operational_impact") or "",
        "is_temporary": review_value("is_temporary") or "",
        "recommendations": review_value("recommendations") or "",
        "signals_json": _signals_for_response(review_value("signals_json")),
        "content_md": review_value("content_md") or "",
        "content_raw": review_value("content_raw") or "",
        "images_json": review_value("images_json") or [],
        "metadata": review_value("metadata") or {},
        "category_id": review_value("category_id"),
        "ai_category_id": review_value("ai_category_id"),
        "ai_category_conf": (
            float(review_value("ai_category_conf")) if review_value("ai_category_conf") is not None else None
        ),
        "ai_category_reason": review_value("ai_category_reason"),
        "status": row["status"],
        "reviewer_id": row["reviewer_id"],
        "review_note": row["review_note"],
        "created_at": row["created_at"].isoformat() if row["created_at"] else None,
        "updated_at": row["updated_at"].isoformat() if row["updated_at"] else None,
        "published_at": row["published_at"].isoformat() if row["published_at"] else None,
        "latest_proposal_revision_id": row["latest_proposal_revision_id"],
        "working_revision_id": row["working_revision_id"],
        "lock_version": row["lock_version"],
        "maintenance_working": working_payload is not None,
        "review_view": "maintenance_working" if working_payload is not None else "entry",
    }


@kbd_router.get("/{kbd_id}/revisions")
async def get_kbd_revisions(request: Request, kbd_id: int) -> dict[str, Any]:
    """获取 Proposal/Expert 历史和当前 runtime active 元数据。"""

    _check_auth(request)
    if _db_manager is None:
        raise HTTPException(status_code=503, detail="数据库未就绪")

    async with _db_manager.async_session_factory() as session:
        entry = await session.get(KbdEntry, kbd_id)
        if entry is None:
            raise HTTPException(status_code=404, detail=f"KBD 条目 {kbd_id} 不存在")

        history_result = await session.execute(
            select(KbdRevision)
            .where(KbdRevision.kbd_entry_id == kbd_id)
            .order_by(KbdRevision.revision_no.desc())
        )
        history = list(history_result.scalars().all())
        revision_by_id = {item.id: item for item in history}
        current_expert, proposal_baseline = select_current_expert_pair(
            history,
            working_revision_id=entry.working_revision_id,
            latest_proposal_revision_id=entry.latest_proposal_revision_id,
        )
        active_result = await session.execute(
            text(
                """
                SELECT r.revision, r.checksum, r.version, r.published_at
                FROM dynamic_resource_active a
                JOIN dynamic_resource_revision r
                  ON r.resource_type = a.resource_type
                 AND r.resource_name = a.resource_name
                 AND r.revision = a.active_revision
                WHERE a.resource_type = 'kbd' AND a.resource_name = :resource_name
                """
            ),
            {"resource_name": str(kbd_id)},
        )
        active = active_result.mappings().first()

    return {
        "kbd_id": kbd_id,
        "latest_proposal_revision_id": entry.latest_proposal_revision_id,
        "working_revision_id": entry.working_revision_id,
        "lock_version": entry.lock_version,
        "expert_signal_edit_summary": summarize_expert_signal_changes(
            proposal_baseline.payload_json if proposal_baseline is not None else {},
            current_expert.payload_json if current_expert is not None else None,
            proposal_revision_id=proposal_baseline.id if proposal_baseline is not None else entry.latest_proposal_revision_id,
            expert_revision_id=current_expert.id if current_expert is not None else None,
        ),
        "history": [
            {
                **(revision_metadata(item) or {}),
                "payload": item.payload_json,
                "generation_metadata": item.generation_metadata or {},
                "diff_from_parent": (
                    diff_revision_payloads(
                        revision_by_id[item.parent_revision_id].payload_json,
                        item.payload_json,
                    )
                    if item.parent_revision_id in revision_by_id
                    else []
                ),
            }
            for item in history
        ],
        "active_resource": (
            {
                "revision": active["revision"],
                "checksum": active["checksum"],
                "version": active["version"],
                "published_at": active["published_at"].isoformat() if active["published_at"] else None,
            }
            if active
            else None
        ),
    }


@kbd_router.post("/{kbd_id}/validate")
async def validate_kbd_candidate(request: Request, kbd_id: int) -> dict[str, Any]:
    """对当前专家工作内容执行无副作用静态 Validation。

    该接口不创建 runtime revision、不切 active。专家可处理问题放入 ``issues``；
    Capability 部署探测属于平台状态，只放入 ``platform_status``，不冒充专家待办。
    """

    _check_auth(request)
    if _db_manager is None:
        raise HTTPException(status_code=503, detail="数据库未就绪")

    async with _db_manager.async_session_factory() as session:
        kbd = await session.get(KbdEntry, kbd_id)
        if kbd is None:
            raise HTTPException(status_code=404, detail=f"KBD 条目 {kbd_id} 不存在")
        if getattr(kbd, "status", None) == "published" and kbd.working_revision_id is not None:
            working_revision = await session.get(KbdRevision, kbd.working_revision_id)
            if (
                working_revision is not None
                and (working_revision.generation_metadata or {}).get("origin") == "admin_maintenance"
            ):
                working_entry = KbdEntry()
                apply_kbd_revision_payload(working_entry, working_revision.payload_json)
                working_entry.id = kbd.id
                working_entry.status = kbd.status
                working_entry.working_revision_id = kbd.working_revision_id
                working_entry.lock_version = kbd.lock_version
                kbd = working_entry

    # ``issues`` 只包含专家能通过编辑当前 KBD 处理的问题。部署探测等平台信息单独
    # 放入 platform_status，避免把不可处理的工程状态伪装成专家审核告警。
    issues: list[dict[str, Any]] = []
    platform_status: list[dict[str, Any]] = []
    for field, label in (
        ("title", "标题"),
        ("problem_description", "问题描述"),
        ("root_cause", "根因"),
        ("solution", "解决方案"),
    ):
        if not str(getattr(kbd, field) or "").strip():
            issues.append(
                {
                    "level": "error",
                    "code": "KBD_REQUIRED_FIELD_MISSING",
                    "location": field,
                    "message": f"{label}不能为空",
                }
            )

    signals_doc = _load_signals_json(kbd.signals_json)
    signals = signals_doc.get("signals") if isinstance(signals_doc, dict) else []
    if not signals:
        issues.append(
            {
                "level": "error",
                "code": "KBD_SIGNALS_MISSING",
                "location": "signals_json",
                "message": "缺少关键信号，请先抽取或人工新增",
            }
        )
    else:
        for issue in expert_editor_issues(signals_doc):
            issues.append(
                {
                    "level": issue["severity"],
                    "code": issue["code"],
                    "location": "关键信号",
                    "message": issue["message"],
                    "action": {"type": "edit_signal_role"},
                }
            )
        try:
            _validate_kbd_publishable_signals_json(signals_doc)
        except jsonschema.ValidationError as exc:
            human_issue = _humanize_signal_validation_error(exc, signals)
            if not any(issue["code"] == human_issue["code"] for issue in issues):
                issues.append(human_issue)

        used_tools: set[str] = set()
        for index, signal in enumerate(signals):
            if not isinstance(signal, dict):
                continue
            tool = str((signal.get("acquire") or {}).get("tool") or "")
            if not tool or tool in used_tools:
                continue
            used_tools.add(tool)
            descriptor = get_capability_descriptor(tool)
            if descriptor is None:
                issues.append(
                    {
                        "level": "error",
                        "code": "CAPABILITY_MISSING",
                        "location": f"signals[{index}].acquire.tool",
                        "message": f"平台代码契约未声明能力 {tool}",
                    }
                )
            elif descriptor["runtime_status"] != "available":
                platform_status.append(
                    {
                        "code": "CAPABILITY_RUNTIME_UNVERIFIED",
                        "capability_id": tool,
                        "status": descriptor["runtime_status"],
                        "message": "当前静态校验只证明参数契约；运行状态请以 Gateway 合并的 Agent 探测结果为准",
                        "expert_action_required": False,
                        "blocks_publish": False,
                    }
                )

        image_type_by_seq: dict[str, str] = {}
        for image in getattr(kbd, "images_json", None) or []:
            if not isinstance(image, dict) or image.get("seq") is None:
                continue
            type_match = re.search(r"^TYPE:\s*(.+?)\s*$", str(image.get("desc") or ""), re.MULTILINE)
            if type_match:
                image_type_by_seq[str(image["seq"])] = type_match.group(1).strip()
        for index, signal in enumerate(signals):
            if not isinstance(signal, dict) or (signal.get("acquire") or {}).get("tool") != "qkv_dialog":
                continue
            source_refs = (signal.get("provenance") or {}).get("source_refs") or []
            referenced_types = {
                image_type_by_seq[match.group(1)]
                for ref in source_refs
                if (match := re.search(r"(?:img(?:age)?[_:#-]?)(\d+)", str(ref), re.IGNORECASE))
                and match.group(1) in image_type_by_seq
            }
            if "任务截图" in referenced_types:
                issues.append(
                    {
                        "level": "warning",
                        "code": "QKV_TASK_PREFERRED",
                        "location": f"signals[{index}].acquire.tool",
                        "message": "该信号引用任务截图：请优先改为 qkv_task；只有确认不存在任务记录时才保留 qkv_dialog",
                        "action": {"type": "edit_signal", "signal_id": signal.get("id"), "suggested_tool": "qkv_task"},
                    }
                )

    if not kbd.category_id and not kbd.ai_category_id:
        issues.append(
            {
                "level": "error",
                "code": "KBD_CATEGORY_MISSING",
                "location": "category_id",
                "message": "缺少人工分类和 AI 建议分类",
            }
        )

    error_count = sum(issue["level"] == "error" for issue in issues)
    warning_count = sum(issue["level"] == "warning" for issue in issues)
    return {
        "kbd_id": kbd_id,
        "working_revision_id": kbd.working_revision_id,
        "lock_version": kbd.lock_version,
        "status": "error" if error_count else "warning" if warning_count else "ok",
        "publishable": error_count == 0,
        "runtime_verified": not platform_status,
        "error_count": error_count,
        "warning_count": warning_count,
        "issues": issues,
        "platform_status": platform_status,
    }


# ─────────────────────────────────────────────────────────────────────────────
# KBD 条目拒绝接口
# ─────────────────────────────────────────────────────────────────────────────


class KbdRejectRequest(BaseModel):
    """KBD 条目拒绝请求"""

    reviewer_id: int = Field(..., description="审核人 ID")
    review_note: str = Field(..., min_length=1, max_length=500, description="拒绝原因（必填）")


@kbd_router.patch("/{kbd_id}/reject")
async def reject_kbd_entry(request: Request, kbd_id: int, body: KbdRejectRequest):
    """拒绝 KBD 条目，更新状态为 rejected"""
    _check_auth(request)

    if _db_manager is None:
        raise HTTPException(status_code=503, detail="数据库未就绪")

    logger.info(event="kbd_reject_request", kbd_id=kbd_id, reviewer_id=body.reviewer_id)

    now = datetime.now(UTC)
    async with _db_manager.async_session_factory() as session:
        result = await session.execute(
            text(
                """
                UPDATE kbd_entry
                SET status = 'rejected',
                    reviewer_id = :reviewer_id,
                    reviewed_at = :reviewed_at,
                    review_note = :review_note
                WHERE id = :id AND status = 'draft'
                RETURNING id, status
                """
            ),
            {
                "id": kbd_id,
                "reviewer_id": body.reviewer_id,
                "reviewed_at": now,
                "review_note": body.review_note,
            },
        )
        updated = result.mappings().first()
        if not updated:
            raise HTTPException(status_code=404, detail=f"KBD 条目 {kbd_id} 不存在或状态非 draft")
        await session.commit()

    logger.info(event="kbd_rejected", kbd_id=kbd_id)
    return {"success": True, "kbd_id": kbd_id, "status": "rejected"}


class KbdApproveRequest(BaseModel):
    """KBD 条目审核通过请求"""

    reviewer_id: int = Field(..., ge=1, description="审核人 ID")
    review_note: str | None = Field(None, max_length=500, description="审核备注（可选）")
    category_id: str | None = Field(
        None,
        description='人工确认的分类 code（可选，如“虚拟机-017”）。来自审核详情弹窗的“确认分类”下拉框，'
        '发布时优先于 AI 自动分类写入 category_id，用于校正分类、避免孤儿 KBD',
    )
    lock_version: int | None = Field(None, ge=0, description="可选专家工作稿版本；不匹配时返回 409")


class KbdApproveResponse(BaseModel):
    """KBD 条目审核通过响应"""

    success: bool = Field(..., description="操作是否成功")
    kbd_id: int = Field(..., description="KBD 条目 ID")
    status: str = Field(..., description="当前状态")
    embedding_generated: bool = Field(..., description="是否成功生成 embedding")
    published_at: str | None = Field(None, description="发布时间")
    resource_revision: dict | None = Field(default=None, description="动态资源 revision 元数据")
    knowledge_revision: dict | None = Field(default=None, description="本次确认的 Expert revision 元数据")


@kbd_router.post("/{kbd_id}/approve", response_model=KbdApproveResponse)
async def approve_kbd_entry(request: Request, kbd_id: int, body: KbdApproveRequest):
    """审核通过 KBD 条目

    功能清单：
    1. 更新 kbd_entry.status → published
    2. 触发 embedding 生成（调用 embedding API 对 content_md 生成向量）
    3. 使用 jieba 生成中文 token，并构建 PostgreSQL tsvector
    4. 设置 published_at = NOW()
    5. 记录 reviewer_id

    响应体示例：
    ```json
    {
      "success": true,
      "kbd_id": 123,
      "status": "published",
      "embedding_generated": true,
      "published_at": "2026-04-02T10:30:00Z"
    }
    ```
    """
    _check_auth(request)

    if _db_manager is None:
        raise HTTPException(status_code=503, detail="数据库未就绪")

    logger.info(
        event="kbd_approve_request",
        kbd_id=kbd_id,
        reviewer_id=body.reviewer_id,
    )

    # 1. 查询 kbd_entry（短事务，快速释放连接）
    async with _db_manager.async_session_factory() as session:
        result = await session.execute(
            text(
                "SELECT id, title, content_md, content_raw, problem_description, alert_info, root_cause, "
                "status, published_at, embedding, signals_json, category_id, ai_category_id, lock_version "
                "FROM kbd_entry WHERE id = :id"
            ),
            {"id": kbd_id},
        )
        row = result.mappings().first()

        if not row:
            raise HTTPException(status_code=404, detail=f"KBD 条目 {kbd_id} 不存在")

        current_status = row["status"]
        if current_status == "published":
            resource_revision = await _publish_kbd_revision(session, kbd_id, get_current_trace_id())
            await session.commit()
            # 已发布，无需重复处理
            return KbdApproveResponse(
                success=True,
                kbd_id=kbd_id,
                status="published",
                embedding_generated=row["embedding"] is not None,
                published_at=row["published_at"].isoformat() if row["published_at"] else None,
                resource_revision=resource_revision,
                knowledge_revision=None,
            )

        source_lock_version = int(row.get("lock_version") or 0)
        if body.lock_version is not None and body.lock_version != source_lock_version:
            raise HTTPException(
                status_code=409,
                detail={
                    "code": "KBD_EDIT_CONFLICT",
                    "message": "该 KBD 已被其他编辑更新，请刷新后重新审核",
                    "expected_lock_version": body.lock_version,
                    "current_lock_version": source_lock_version,
                },
            )

        content_md = row["content_md"]
        if not content_md:
            raise HTTPException(
                status_code=400,
                detail=f"KBD 条目 {kbd_id} 缺少 content_md，无法生成 embedding",
            )

        # ADR-1 + ADR-3：审核前置门
        # 门 1：signals_json 非空（无关键信号 → CDD 不可执行）
        # 直接切 v2 列形态（RFC §7）：signals_json 现为 {schema_version, signals} 对象，
        # 需先解包为信号列表再判空。
        _signals_doc = _load_signals_json(row["signals_json"]) or {}
        _raw_signals = _signals_doc.get("signals", [])
        if not _raw_signals:
            raise HTTPException(
                status_code=422,
                detail=f"KBD 条目 {kbd_id} 缺少关键信号（signals_json 为空），请先调用 /extract-signals 抽取后再审核",
            )
        try:
            _signals_doc = _prepare_expert_publish_signals(_signals_doc)
        except jsonschema.ValidationError as exc:
            _raise_signal_validation_error(exc, _signals_doc.get("signals") or [])
        # 门 1.5：至少含 1 条消费者(backend)信号，否则 CDD 无法执行差异消除（§9）
        # v2 原生判定：acquire.tool 以 qfk 开头 或 provenance.category==backend
        _has_consumer = any(
            isinstance(s, dict)
            and (
                (s.get("acquire") or {}).get("tool", "").startswith("qfk")
                or (s.get("provenance") or {}).get("category") == "backend"
            )
            for s in _raw_signals
        )
        if not _has_consumer:
            raise HTTPException(
                status_code=422,
                detail=f"KBD 条目 {kbd_id} 缺少消费者(backend)信号，CDD 无法进行差异诊断消除，请补充至少 1 条 QFK 消费者信号后再审核",
            )
        # 门 2：category_id 与 ai_category_id 同步（根治孤儿 KBD）
        # 优先采用人工确认的分类（审核详情弹窗“确认分类”下拉框），
        # fallback 到 DB 已有值；保证发布后 category_id 一定有值。
        effective_category_id = body.category_id or row["category_id"] or row["ai_category_id"]
        if not effective_category_id:
            raise HTTPException(
                status_code=422,
                detail=f"KBD 条目 {kbd_id} 缺少分类（category_id 与 ai_category_id 均为空），请先抽取分类后再审核",
            )
        embedding_text = build_kbd_embedding_text(
            title=row["title"],
            problem_description=row["problem_description"],
            alert_info=row["alert_info"],
            root_cause=row["root_cause"],
            fallback_text=row["content_raw"] or content_md,
        )

    # 2. 生成 embedding（事务外调用，避免长时间占用连接）
    embedding_generated = False
    embedding_vector: list[float] | None = None
    embedding_content_hash = hashlib.sha256(embedding_text.encode("utf-8")).hexdigest()
    if _embedding_service:
        try:
            embedding_vector = await _embedding_service.embed_single(embedding_text)
            embedding_generated = True

            actual_dim = len(embedding_vector)
            logger.info(
                event="kbd_embedding_generated",
                kbd_id=kbd_id,
                vector_dim=actual_dim,
            )
        except Exception as exc:
            logger.warning(
                event="kbd_embedding_failed",
                kbd_id=kbd_id,
                error=str(exc),
                message="embedding 生成失败，将继续更新状态，后续可手动重试",
            )

    # 3. 更新 kbd_entry 状态（短事务）
    now = datetime.now(UTC)
    current_content_raw = strip_markdown(content_md or "")
    tsv_text = segment(f"{row['title']} {current_content_raw}")
    async with _db_manager.async_session_factory() as session:
        # 构建 UPDATE SQL（embedding 使用 pgvector 格式）
        if embedding_vector:
            # 将向量列表转换为 PostgreSQL vector 格式字符串
            vector_str = "[" + ",".join(str(v) for v in embedding_vector) + "]"
            update_sql = text(
                """
                UPDATE kbd_entry
                SET status = 'published',
                    published_at = :published_at,
                    reviewer_id = :reviewer_id,
                    reviewed_at = :reviewed_at,
                    review_note = COALESCE(:review_note, review_note),
                    content_raw = :content_raw,
                    category_id = :category_id,
                    signals_json = CAST(:signals_json AS jsonb),
                    embedding = CAST(:embedding AS vector),
                    embedding_model = :embedding_model,
                    embedding_content_hash = :embedding_content_hash,
                    embedding_updated_at = :embedding_updated_at,
                    tsv = to_tsvector('simple', :tsv_text)
                WHERE id = :id AND lock_version = :expected_lock_version AND status <> 'published'
                RETURNING id, status, embedding, published_at
                """
            )
            params = {
                "id": kbd_id,
                "expected_lock_version": source_lock_version,
                "published_at": now,
                "reviewer_id": body.reviewer_id,
                "reviewed_at": now,
                "review_note": body.review_note,
                "content_raw": current_content_raw,
                "category_id": effective_category_id,
                "signals_json": json.dumps(_signals_doc, ensure_ascii=False),
                "embedding": vector_str,
                "embedding_model": _embedding_service.model_name,
                "embedding_content_hash": embedding_content_hash,
                "embedding_updated_at": now,
                "tsv_text": tsv_text,
            }
        else:
            # 生成失败时必须清空向量，避免旧内容向量或伪向量继续参与召回
            update_sql = text(
                """
                UPDATE kbd_entry
                SET status = 'published',
                    published_at = :published_at,
                    reviewer_id = :reviewer_id,
                    reviewed_at = :reviewed_at,
                    review_note = COALESCE(:review_note, review_note),
                    content_raw = :content_raw,
                    category_id = :category_id,
                    signals_json = CAST(:signals_json AS jsonb),
                    embedding = NULL,
                    embedding_model = NULL,
                    embedding_content_hash = NULL,
                    embedding_updated_at = NULL,
                    tsv = to_tsvector('simple', :tsv_text)
                WHERE id = :id AND lock_version = :expected_lock_version AND status <> 'published'
                RETURNING id, status, embedding, published_at
                """
            )
            params = {
                "id": kbd_id,
                "expected_lock_version": source_lock_version,
                "published_at": now,
                "reviewer_id": body.reviewer_id,
                "reviewed_at": now,
                "review_note": body.review_note,
                "content_raw": current_content_raw,
                "category_id": effective_category_id,
                "signals_json": json.dumps(_signals_doc, ensure_ascii=False),
                "tsv_text": tsv_text,
            }

        result = await session.execute(update_sql, params)
        updated = result.mappings().first()

        if not updated:
            raise HTTPException(
                status_code=409,
                detail={
                    "code": "KBD_EDIT_CONFLICT",
                    "message": "生成发布数据期间 KBD 已发生变化，未发布旧版本；请刷新后重新审核",
                    "expected_lock_version": source_lock_version,
                },
            )

        # 通过门禁后才冻结 Expert Revision；现有 body reviewer_id 尚不是可信认证身份，
        # helper 会诚实记录 identity_status，避免伪造 Gold actor。
        expert_revision = await _freeze_approved_expert_revision(
            session,
            kbd_id=kbd_id,
            reviewer_id=body.reviewer_id,
            review_note=body.review_note,
            trace_id=get_current_trace_id(),
        )
        resource_revision = await _publish_kbd_revision(session, kbd_id, get_current_trace_id())
        published_entry = await session.get(KbdEntry, kbd_id)
        if published_entry is not None:
            published_entry.working_revision_id = None
        await session.commit()

    logger.info(
        event="kbd_approved",
        kbd_id=kbd_id,
        reviewer_id=body.reviewer_id,
        embedding_generated=embedding_generated,
    )

    return KbdApproveResponse(
        success=True,
        kbd_id=kbd_id,
        status=updated["status"],
        embedding_generated=updated["embedding"] is not None,
        published_at=updated["published_at"].isoformat() if updated["published_at"] else None,
        resource_revision=resource_revision,
        knowledge_revision=revision_metadata(expert_revision),
    )


# ─────────────────────────────────────────────────────────────────────────────
# SOP 文档审核接口（sop_document 单表，tree_json 合并存储）
# ─────────────────────────────────────────────────────────────────────────────


class SopApproveRequest(BaseModel):
    """SOP 文档审核通过请求"""

    reviewer_id: int = Field(..., ge=1, description="审核人 ID")
    review_note: str | None = Field(None, max_length=500, description="审核备注（可选）")


class SopApproveResponse(BaseModel):
    """SOP 文档审核通过响应"""

    success: bool = Field(..., description="操作是否成功")
    document_id: int = Field(..., description="SOP 文档 ID")
    status: str = Field(..., description="当前状态")
    tree_generated: bool = Field(..., description="是否成功生成 SOP 决策树")
    tree_leaf_count: int | None = Field(None, description="决策树叶节点数量")
    tree_validation_status: str | None = Field(None, description="决策树校验状态（valid/warnings/error）")
    variable_count: int = Field(0, description="提取的变量数量（T-AGT-24）")
    warnings: list[str] = Field(default_factory=list, description="审核警告列表（兼容旧格式，含 orphan 变量等）")
    validation_issues: list[ValidationIssue] = Field(
        default_factory=list,
        description="决策树校验问题列表（含 line_number，供前端按行定位）",
    )
    published_at: str | None = Field(None, description="发布时间")
    resource_revision: dict | None = Field(default=None, description="动态资源 revision 元数据")


async def validate_variable_schema_dependencies(session, variable_schema: list[dict]) -> None:
    """验证变量 Schema 中声明的工具/技能可用性。

    如果存在未启用或未注册的工具/技能，抛出 422 错误。
    """
    required_tools = set()
    required_skills = set()
    for var in variable_schema:
        strat_raw = var.get("acquisition_strategy")
        if not strat_raw:
            continue
        parsed = parse_strategy(strat_raw)
        tool_name = var.get("acquisition_tool") or parsed.acquisition_tool
        if parsed.strategy == "tool_call" and tool_name:
            required_tools.add(tool_name)
        elif parsed.strategy == "skill_call" and tool_name:
            required_skills.add(tool_name)

    missing_tools = set()
    if required_tools:
        stmt = select(ToolDefinitionORM.tool_name).where(
            ToolDefinitionORM.tool_name.in_(list(required_tools)),
            ToolDefinitionORM.is_active.is_(True),
        )
        res = await session.execute(stmt)
        active_tools = set(res.scalars().all())
        missing_tools = required_tools - active_tools

    missing_skills = set()
    if required_skills:
        stmt = select(SkillDefinitionORM.skill_name).where(
            SkillDefinitionORM.skill_name.in_(list(required_skills)),
            SkillDefinitionORM.is_active.is_(True),
        )
        res = await session.execute(stmt)
        active_skills = set(res.scalars().all())
        missing_skills = required_skills - active_skills

    if missing_tools or missing_skills:
        missing_details = []
        issues = []
        if missing_tools:
            missing_details.append(f"工具：{', '.join(sorted(missing_tools))}")
            for t in sorted(missing_tools):
                issues.append(
                    {
                        "level": "error",
                        "location": "变量声明",
                        "line_number": None,
                        "message": f"依赖了未注册或未启用的工具：'{t}'，请先创建或启用它。",
                    }
                )
        if missing_skills:
            missing_details.append(f"技能：{', '.join(sorted(missing_skills))}")
            for s in sorted(missing_skills):
                issues.append(
                    {
                        "level": "error",
                        "location": "变量声明",
                        "line_number": None,
                        "message": f"依赖了未注册或未启用的技能：'{s}'，请先创建或启用它。",
                    }
                )

        detail_msg = "、".join(missing_details)
        raise HTTPException(
            status_code=422,
            detail={
                "error": "missing_dependencies",
                "message": f"SOP 依赖了未注册或未启用的 {detail_msg}，请先创建或启用它们。",
                "missing_tools": list(missing_tools),
                "missing_skills": list(missing_skills),
                "validation_issues": issues,
            },
        )


@sop_router.post("/{document_id}/approve", response_model=SopApproveResponse)
async def approve_sop_document(request: Request, document_id: int, body: SopApproveRequest):
    """审核通过 SOP 文档

    功能清单：
    1. 更新 sop_document.status → published
    2. 解析 content_md 生成 SOP 决策树（SOPNode JSON），写入 sop_document.tree_json
    3. 设置 published_at = NOW()
    4. 记录 reviewer_id

    两段式事务设计（解析树不持有 DB 连接）：
      - 短事务1：查询 document（验证存在）
      - 无事务：解析 SOP Markdown 生成决策树（无 IO 操作，但解析可能耗时）
      - 短事务2：UPDATE sop_document（状态 + tree_json）

    响应体示例：
    ```json
    {
      "success": true,
      "document_id": 1,
      "status": "published",
      "published_at": "2026-04-02T10:30:00Z"
    }
    ```
    """
    _check_auth(request)

    if _db_manager is None:
        raise HTTPException(status_code=503, detail="数据库未就绪")

    logger.info(
        event="sop_approve_request",
        document_id=document_id,
        reviewer_id=body.reviewer_id,
    )

    try:
        # ── 短事务1：查询验证（快速释放连接）────────────────────────────────────
        async with _db_manager.async_session_factory() as session:
            result = await session.execute(select(SopDocument).where(SopDocument.id == document_id))
            sop_doc = result.scalar_one_or_none()

            if not sop_doc:
                raise HTTPException(status_code=404, detail=f"SOP 文档 {document_id} 不存在")

            if sop_doc.status == "published":
                # 已发布，直接返回当前 tree 信息
                # 获取 variable_schema（若存在）
                var_schema_raw = sop_doc.variable_schema or []
                resource_revision = await _publish_sop_revision(session, document_id, get_current_trace_id())
                await session.commit()
                return SopApproveResponse(
                    success=True,
                    document_id=document_id,
                    status="published",
                    tree_generated=sop_doc.tree_json is not None,
                    tree_leaf_count=sop_doc.tree_leaf_count if sop_doc.tree_json is not None else None,
                    tree_validation_status=sop_doc.tree_validation_status if sop_doc.tree_json is not None else None,
                    variable_count=len(var_schema_raw),
                    warnings=[],  # 已发布不再返回历史警告
                    published_at=sop_doc.published_at.isoformat() if sop_doc.published_at else None,
                    resource_revision=resource_revision,
                )

            content_md = sop_doc.content_md
            # 获取旧的 variable_schema（用于三路合并）
            old_variable_schema: list[dict] = sop_doc.variable_schema or []

        # ── 无事务：解析 SOP 决策树（不持有 DB 连接）──────────────────────────
        now = datetime.now(UTC)
        tree_generated = False
        tree_leaf_count = 0
        tree_validation_status: str | None = None
        tool_contract_issues: list[ValidationIssue] = []

        if content_md:
            parse_result = parse_sop_markdown(content_md)
            if not parse_result.has_error and parse_result.root_nodes:
                tree_generated = True
                root = parse_result.root_nodes[0]
                tree_leaf_count = len(_collect_leaves(root))
                tool_contract_issues = validate_sop_tool_contract(root)
                parse_result.issues.extend(tool_contract_issues)
                # 判断是否有 warning
                has_warnings = any(i.level == "warning" for i in parse_result.issues)
                tree_validation_status = "warnings" if has_warnings else "valid"
                logger.info(
                    event="sop_tree_parsed",
                    document_id=document_id,
                    leaf_count=tree_leaf_count,
                    warning_count=len([i for i in parse_result.issues if i.level == "warning"]),
                    tool_contract_warning_count=len(tool_contract_issues),
                )
            else:
                tree_validation_status = "error"
                error_issues = [i for i in parse_result.issues if i.level == "error"]
                logger.warning(
                    event="sop_tree_parse_failed",
                    document_id=document_id,
                    error_count=len(error_issues),
                    errors=[e.message for e in error_issues[:3]],
                )
        else:
            parse_result = None
            logger.warning(
                event="sop_tree_no_content",
                document_id=document_id,
                message="SOP 文档没有 content_md，无法生成决策树",
            )

        # ── 变量提取 + 双向校验（T-AGT-24）────────────────────────────────────────
        variable_defs: list[dict] = []
        undeclared_errors: list[str] = []
        orphan_warnings: list[str] = []
        warnings: list[str] = []
        deprecated_vars: list[str] = []  # 三路合并中标记 deprecated 的变量名

        if content_md:
            # 提取变量（传入解析后的决策树，扫描节点中的变量占位符）
            tree_for_var = (
                parse_result.root_nodes[0]
                if (parse_result and not parse_result.has_error and parse_result.root_nodes)
                else None
            )
            new_variable_defs, undeclared_errors, orphan_warnings = extract_sop_variables(content_md, tree_for_var)

            # Undeclared = Error（阻断 approve）
            if undeclared_errors:
                logger.warning(
                    event="sop_undeclared_variables",
                    document_id=document_id,
                    undeclared_vars=undeclared_errors,
                )
                raise HTTPException(
                    status_code=422,
                    detail={
                        "error": "undeclared_variables",
                        "message": f"SOP 正文使用了未声明的变量：{undeclared_errors}，请在 ## 变量 章节中声明",
                        "undeclared": undeclared_errors,
                    },
                )

            # Orphan = Warning（写入响应）
            for var_name in orphan_warnings:
                warnings.append(f"变量 '{var_name}' 已在 ## 变量 章节声明但未在正文中使用")

            # 三路合并（T-AGT-26）：合并新旧 variable_schema
            if old_variable_schema:
                variable_defs, deprecated_vars = merge_variable_schema(old_variable_schema, new_variable_defs)
                # Deprecated 变量告警（写入响应）
                for var_name in deprecated_vars:
                    warnings.append(f"变量 '{var_name}' 已从新版 SOP 中移除，标记为 deprecated")
                logger.info(
                    event="sop_variable_merge",
                    document_id=document_id,
                    old_count=len(old_variable_schema),
                    new_count=len(new_variable_defs),
                    merged_count=len(variable_defs),
                    deprecated_count=len(deprecated_vars),
                )
            else:
                # 无旧版 schema，直接使用新版
                variable_defs = new_variable_defs

            # 合并决策树解析警告（成功时）
            if parse_result and not parse_result.has_error:
                for w in parse_result.issues:
                    if w.level == "warning":
                        warnings.append(f"[决策树] {w.location}: {w.message}")

            # 当决策树解析失败时，将错误信息追加到 warnings
            if parse_result and parse_result.has_error:
                for e in parse_result.issues:
                    if e.level == "error":
                        warnings.append(f"[决策树解析失败] {e.location}: {e.message}")

            logger.info(
                event="sop_variables_extracted",
                document_id=document_id,
                variable_count=len(variable_defs),
                orphan_count=len(orphan_warnings),
                deprecated_count=len(deprecated_vars),
            )

        # ── 短事务2：UPDATE sop_document（状态 + tree_json）────────────────────
        async with _db_manager.async_session_factory() as session:
            # 校验工具与技能依赖（T-AGT-28）
            if parse_result and not parse_result.has_error and parse_result.root_nodes:
                await validate_variable_schema_dependencies(session, variable_defs)

            # 决策树解析成功才设置 status = published，失败则保持 draft 并记录错误
            if parse_result and not parse_result.has_error and parse_result.root_nodes:
                # 解析成功：更新为 published + 写入决策树
                root = parse_result.root_nodes[0]
                await session.execute(
                    text(
                        """
                        UPDATE sop_document
                        SET status = 'published',
                            published_at = :published_at,
                            reviewer_id = :reviewer_id,
                            reviewed_at = :reviewed_at,
                            review_note = COALESCE(:review_note, review_note),
                            tree_json = CAST(:tree_json AS jsonb),
                            tree_schema_version = :schema_version,
                            tree_leaf_count = :leaf_count,
                            tree_validation_status = :validation_status,
                            tree_validation_issues = CAST(:validation_issues AS jsonb),
                            tree_generator_version = :generator_version,
                            variable_schema = CAST(:variable_schema AS jsonb),
                            updated_at = :updated_at
                        WHERE id = :id
                        """
                    ),
                    {
                        "id": document_id,
                        "published_at": now,
                        "reviewer_id": body.reviewer_id,
                        "reviewed_at": now,
                        "review_note": body.review_note,
                        "tree_json": json.dumps(root.model_dump(), ensure_ascii=False),
                        "schema_version": "sop-tree-v1",
                        "leaf_count": tree_leaf_count,
                        "validation_status": tree_validation_status,
                        "validation_issues": json.dumps(
                            [i.model_dump() for i in parse_result.issues], ensure_ascii=False
                        ),
                        "generator_version": "sop-parser-v1",
                        "variable_schema": json.dumps(variable_defs, ensure_ascii=False) if variable_defs else None,
                        "updated_at": now,
                    },
                )
                logger.info(
                    event="sop_tree_written",
                    document_id=document_id,
                    leaf_count=tree_leaf_count,
                    variable_count=len(variable_defs),
                )
            elif parse_result and parse_result.has_error:
                # 解析失败：保持 draft + 记录错误信息（不阻断，返回 422 让用户修复）
                await session.execute(
                    text(
                        """
                        UPDATE sop_document
                        SET tree_validation_status = 'error',
                            tree_validation_issues = CAST(:validation_issues AS jsonb),
                            updated_at = :updated_at
                        WHERE id = :id
                        """
                    ),
                    {
                        "id": document_id,
                        "validation_issues": json.dumps(
                            [i.model_dump() for i in parse_result.issues], ensure_ascii=False
                        ),
                        "updated_at": now,
                    },
                )
                # 构建错误详情返回给前端
                error_issues = [i for i in parse_result.issues if i.level == "error"]
                raise HTTPException(
                    status_code=422,
                    detail={
                        "error": "tree_parse_failed",
                        "message": f"SOP 决策树解析失败，共 {len(error_issues)} 个错误，请修复后重新发布",
                        "validation_issues": [i.model_dump() for i in parse_result.issues],
                    },
                )
            else:
                # 无 content_md 或解析结果为空：保持 draft 状态，记录原因
                await session.execute(
                    text(
                        """
                        UPDATE sop_document
                        SET review_note = COALESCE(:review_note, review_note),
                            updated_at = :updated_at
                        WHERE id = :id
                        """
                    ),
                    {
                        "id": document_id,
                        "review_note": body.review_note or "无法生成决策树：文档内容为空或无有效标题",
                        "updated_at": now,
                    },
                )
                logger.warning(
                    event="sop_approve_no_tree",
                    document_id=document_id,
                    message="无法生成决策树，保持 draft 状态",
                )
                raise HTTPException(
                    status_code=422,
                    detail={
                        "error": "no_tree_generated",
                        "message": "无法生成决策树：文档内容为空或无有效标题",
                    },
                )

            resource_revision = await _publish_sop_revision(session, document_id, get_current_trace_id())
            await session.commit()

    except HTTPException:
        raise
    except Exception as exc:
        logger.exception(
            event="sop_approve_unexpected_error",
            message="发布 SOP 文档时发生未预期异常",
            document_id=document_id,
            error_type=type(exc).__name__,
        )
        raise HTTPException(
            status_code=500,
            detail="发布 SOP 文档失败，请联系管理员或查看服务日志",
        ) from exc

    logger.info(
        event="sop_approved",
        document_id=document_id,
        reviewer_id=body.reviewer_id,
        tree_generated=tree_generated,
        variable_count=len(variable_defs),
        tool_contract_warning_count=len(tool_contract_issues),
    )

    # 构建 validation_issues：合并 parse_result 中的所有 issues（含 line_number）
    validation_issues: list[dict] = []
    if parse_result:
        for issue in parse_result.issues:
            validation_issues.append(issue.model_dump())

    return SopApproveResponse(
        success=True,
        document_id=document_id,
        status="published",
        tree_generated=tree_generated,
        tree_leaf_count=tree_leaf_count if tree_generated else None,
        tree_validation_status=tree_validation_status,
        variable_count=len(variable_defs),
        warnings=warnings,
        validation_issues=validation_issues,
        published_at=now.isoformat(),
        resource_revision=resource_revision,
    )


# ─────────────────────────────────────────────────────────────────────────────
# SOP 文档单条详情查询（含 content_md）
# ─────────────────────────────────────────────────────────────────────────────


@sop_router.get("/{document_id}")
async def get_sop_document(request: Request, document_id: int):
    """获取单个 SOP 文档详情（含 content_md 正文和 variable_schema）"""
    _check_auth(request)

    if _db_manager is None:
        raise HTTPException(status_code=503, detail="数据库未就绪")

    async with _db_manager.async_session_factory() as session:
        row = (
            (
                await session.execute(
                    text(
                        """
                SELECT id, source_id, category_id, title, content_md, status,
                       reviewer_id, reviewed_at, published_at, created_at, updated_at,
                       tree_leaf_count, (tree_json IS NOT NULL) AS has_tree,
                       tree_validation_status, tree_validation_issues,
                       variable_schema
                FROM sop_document WHERE id = :id
                """
                    ),
                    {"id": document_id},
                )
            )
            .mappings()
            .first()
        )

    if not row:
        raise HTTPException(status_code=404, detail=f"SOP 文档 {document_id} 不存在")

    return {
        "id": row["id"],
        "source_id": row["source_id"],
        "category_id": row["category_id"],
        "title": row["title"],
        "content_md": row["content_md"],
        "status": row["status"],
        "tree_leaf_count": row["tree_leaf_count"],
        "has_tree": row["has_tree"],
        "tree_validation_status": row["tree_validation_status"],
        "tree_validation_issues": row["tree_validation_issues"] or [],
        "variable_schema": row["variable_schema"] or [],
        "reviewer_id": row["reviewer_id"],
        "reviewed_at": row["reviewed_at"].isoformat() if row["reviewed_at"] else None,
        "published_at": row["published_at"].isoformat() if row["published_at"] else None,
        "created_at": row["created_at"].isoformat() if row["created_at"] else None,
        "updated_at": row["updated_at"].isoformat() if row["updated_at"] else None,
    }


# ─────────────────────────────────────────────────────────────────────────────
# SOP 决策树查询接口（供管理前端渲染树结构）
# ─────────────────────────────────────────────────────────────────────────────


@sop_router.get("/{document_id}/tree")
async def get_sop_tree(request: Request, document_id: int):
    """获取 SOP 文档的决策树 JSON（tree_json 字段）。

    返回值为 SOPNode.model_dump() 格式，树根节点对象。
    文档不存在或决策树尚未生成时返回 404。
    """
    _check_auth(request)

    if _db_manager is None:
        raise HTTPException(status_code=503, detail="数据库未就绪")

    async with _db_manager.async_session_factory() as session:
        row = (
            (
                await session.execute(
                    text(
                        """
                SELECT id, title, tree_json, tree_validation_status, tree_leaf_count
                FROM sop_document WHERE id = :id
                """
                    ),
                    {"id": document_id},
                )
            )
            .mappings()
            .first()
        )

    if not row:
        raise HTTPException(status_code=404, detail=f"SOP 文档 {document_id} 不存在")

    if row["tree_json"] is None:
        raise HTTPException(
            status_code=404,
            detail=f"SOP 文档 {document_id} 的决策树尚未生成，请先发布文档",
        )

    return {
        "document_id": row["id"],
        "title": row["title"],
        "tree_validation_status": row["tree_validation_status"],
        "tree_leaf_count": row["tree_leaf_count"],
        "tree": row["tree_json"],
    }


# ─────────────────────────────────────────────────────────────────────────────
# SOP 文档列表查询接口
# ─────────────────────────────────────────────────────────────────────────────


@sop_router.get("")
async def list_sop_documents(
    request: Request,
    page: int = 1,
    page_size: int = 20,
    status: str | None = None,
    category_id: str | None = None,
):
    """查询 SOP 文档列表（分页 + 状态/分类过滤）

    Returns:
        { documents: [...], total, page, page_size }
    """
    _check_auth(request)

    if _db_manager is None:
        raise HTTPException(status_code=503, detail="数据库未就绪")

    page_size = min(max(page_size, 1), 100)
    page = max(page, 1)
    offset = (page - 1) * page_size

    async with _db_manager.async_session_factory() as session:
        where_clauses = []
        params: dict = {"limit": page_size, "offset": offset}

        if status:
            where_clauses.append("status = :status")
            params["status"] = status
        if category_id:
            where_clauses.append("category_id = :category_id")
            params["category_id"] = category_id

        where_sql = ("WHERE " + " AND ".join(where_clauses)) if where_clauses else ""

        count_sql = text(f"SELECT COUNT(*) FROM sop_document {where_sql}")  # noqa: S608
        count_result = await session.execute(count_sql, params)
        total = count_result.scalar() or 0

        data_sql = text(  # noqa: S608
            f"""
            SELECT id, source_id, category_id, title, status,
                   reviewer_id, reviewed_at, published_at, created_at, updated_at, hit_count,
                   tree_leaf_count, (tree_json IS NOT NULL) AS has_tree, tree_validation_status
            FROM sop_document
            {where_sql}
            ORDER BY created_at DESC, id DESC
            LIMIT :limit OFFSET :offset
            """
        )
        result = await session.execute(data_sql, params)
        rows = result.mappings().all()

    documents = [
        {
            "id": row["id"],
            "source_id": row["source_id"],
            "category_id": row["category_id"],
            "title": row["title"],
            "status": row["status"],
            "tree_leaf_count": row["tree_leaf_count"],
            "has_tree": row["has_tree"],
            "tree_validation_status": row["tree_validation_status"],
            "reviewer_id": row["reviewer_id"],
            "reviewed_at": row["reviewed_at"].isoformat() if row["reviewed_at"] else None,
            "published_at": row["published_at"].isoformat() if row["published_at"] else None,
            "created_at": row["created_at"].isoformat() if row["created_at"] else None,
            "updated_at": row["updated_at"].isoformat() if row["updated_at"] else None,
            "hit_count": row.get("hit_count", 0),
        }
        for row in rows
    ]

    return {
        "documents": documents,
        "total": total,
        "page": page,
        "page_size": page_size,
    }


# ─────────────────────────────────────────────────────────────────────────────
# SOP 文档状态更新接口（下线/归档）
# ─────────────────────────────────────────────────────────────────────────────


class SopStatusUpdateRequest(BaseModel):
    """SOP 文档状态/信息更新请求"""

    status: str | None = Field(None, description="目标状态：archived 等")
    title: str | None = Field(None, max_length=500, description="新标题（可选）")
    category_id: str | None = Field(None, max_length=32, description="新分类 ID（可选，传空字符串清除）")
    content_md: str | None = Field(None, description="更新后的 Markdown 正文（可选，修改后将重新分块）")


@sop_router.patch("/{document_id}")
async def update_sop_status(request: Request, document_id: int, body: SopStatusUpdateRequest):
    """更新 SOP 文档状态、标题或分类"""
    _check_auth(request)

    if _db_manager is None:
        raise HTTPException(status_code=503, detail="数据库未就绪")

    if body.status is not None and body.status not in SopDocument.VALID_STATUSES:
        raise HTTPException(
            status_code=400,
            detail=f"非法状态: {body.status}，合法值: {list(SopDocument.VALID_STATUSES)}",
        )

    if body.status is None and body.title is None and body.category_id is None and body.content_md is None:
        raise HTTPException(status_code=400, detail="至少需要提供一个更新字段")

    rechunked = False
    downgraded_to_draft = False

    async with _db_manager.async_session_factory() as session:
        result = await session.execute(select(SopDocument).where(SopDocument.id == document_id))
        sop_doc = result.scalar_one_or_none()
        if not sop_doc:
            raise HTTPException(status_code=404, detail=f"SOP 文档 {document_id} 不存在")

        if body.status is not None:
            sop_doc.status = body.status
        if body.title is not None:
            sop_doc.title = body.title
        if body.category_id is not None:
            # 传空字符串表示清除分类
            sop_doc.category_id = body.category_id or None

        if body.content_md is not None:
            sop_doc.content_md = body.content_md
            # 内容变更后清空决策树（需重新发布生成）
            sop_doc.tree_json = None
            sop_doc.tree_validation_status = None
            rechunked = True
            # 内容变更后若已发布则降级为草稿
            if sop_doc.status == "published" and body.status is None:
                sop_doc.status = "draft"
                downgraded_to_draft = True

        await session.commit()

    logger.info(
        event="sop_updated",
        document_id=document_id,
        new_status=sop_doc.status,
        new_title=body.title,
        content_updated=rechunked,
        downgraded=downgraded_to_draft,
    )
    resp = {"success": True, "document_id": document_id, "status": sop_doc.status}
    if downgraded_to_draft:
        resp["message"] = "内容已更新，决策树已清空，文档已降级为草稿，请重新发布"
    return resp


# ─────────────────────────────────────────────────────────────────────────────
# SOP 变量 Schema 编辑接口（不触发 re-approve）
# ─────────────────────────────────────────────────────────────────────────────


class SopVariableSchemaUpdateRequest(BaseModel):
    """SOP 变量 Schema 编辑请求（T-AGT-28）

    仅更新指定变量的可编辑字段，不触发 re-approve（不清空 tree_json）。
    可编辑字段：display_name、description、acquisition_strategy、acquisition_prompt、
                validation_pattern、acquisition_tool、default_value、depends_on、
                output_path、fallback_strategy、acquisition_args_template、expression
    """

    variables: list[dict] = Field(
        ...,
        min_length=1,
        description="需要更新的变量列表，每项必须包含 name 字段",
    )


@sop_router.patch("/{document_id}/variable-schema")
async def update_sop_variable_schema(request: Request, document_id: int, body: SopVariableSchemaUpdateRequest):
    """更新 SOP 变量 Schema 的可编辑字段（不触发 re-approve）

    功能：
    1. 仅更新指定变量的可编辑字段（display_name、description 等）
    2. 不触发 re-approve（保持 status 和 tree_json 不变）
    3. 三路合并兼容：下次 approve 时保留人工编辑字段

    Args:
        document_id: SOP 文档 ID
        body.variables: 需要更新的变量列表，每项格式：
            {
              "name": "vm_name",                     # 必填，变量名（用于匹配）
              "display_name": "虚拟机名称",          # 可选
              "description": "需要操作的虚拟机",      # 可选
              "acquisition_strategy": "user_confirm",# 可选
              "acquisition_prompt": "请确认虚拟机",   # 可选
              "acquisition_tool": "get_vm_list",     # 可选
              "validation_pattern": "^[a-zA-Z0-9_-]+$", # 可选
              "default_value": "default-vm",         # 可选
              "depends_on": ["node_ip"],             # 可选
              "output_path": "stdout",               # 可选
              "fallback_strategy": "user_input",     # 可选
              "acquisition_args_template": {},       # 可选
              "expression": "contains(alert_type, 'vs') ? false : unknown" # 可选
            }

    Returns:
        { success, document_id, updated, variable_schema }
    """
    _check_auth(request)

    if _db_manager is None:
        raise HTTPException(status_code=503, detail="数据库未就绪")

    logger.info(
        event="sop_variable_schema_update_request", document_id=document_id, variables_count=len(body.variables)
    )

    async with _db_manager.async_session_factory() as session:
        # 1. 查询当前 variable_schema
        result = await session.execute(
            text("SELECT id, status, variable_schema FROM sop_document WHERE id = :id"),
            {"id": document_id},
        )
        row = result.mappings().first()

        if not row:
            raise HTTPException(status_code=404, detail=f"SOP 文档 {document_id} 不存在")

        current_schema: list[dict] = row["variable_schema"] or []

        if not current_schema:
            raise HTTPException(
                status_code=400,
                detail=f"SOP 文档 {document_id} 无 variable_schema，请先 approve 生成",
            )

        # 2. 构建更新后的 schema（保留未更新变量）
        current_by_name = {v["name"]: v for v in current_schema}
        updated_count = 0
        allowed_fields = {
            "display_name",
            "description",
            "acquisition_strategy",
            "acquisition_prompt",
            "acquisition_tool",
            "validation_pattern",
            "default_value",
            "depends_on",
            "output_path",
            "fallback_strategy",
            "acquisition_args",
            "acquisition_args_template",
            "expression",
        }

        for update_var in body.variables:
            var_name = update_var.get("name")
            if not var_name:
                raise HTTPException(
                    status_code=400,
                    detail=f"variables 列表中缺少 name 字段：{update_var}",
                )

            if var_name not in current_by_name:
                raise HTTPException(
                    status_code=400,
                    detail=f"变量 '{var_name}' 不存在于当前 variable_schema 中",
                )

            # 更新允许的字段（仅更新传入的字段）
            current_var = current_by_name[var_name]
            for field, value in update_var.items():
                if field == "name":
                    continue  # name 用于匹配，不可修改
                if field not in allowed_fields:
                    raise HTTPException(
                        status_code=400,
                        detail=f"字段 '{field}' 不允许编辑，可编辑字段：{sorted(allowed_fields)}",
                    )
                # DC-04: validation_pattern 需验证为合法正则，防止写入无效值导致运行时 500
                if field == "validation_pattern" and value:
                    try:
                        re.compile(value)
                    except re.error as exc:
                        raise HTTPException(
                            status_code=400,
                            detail=f"变量 '{var_name}' 的 validation_pattern '{value}' 不是合法正则: {exc}",
                        )
                current_var[field] = value
                updated_count += 1

            # 标记为人工编辑（下次 approve 保留）
            current_var["auto_generated"] = False

        # 校验修改后的整个 schema 依赖
        await validate_variable_schema_dependencies(session, current_schema)

        # 3. 写回数据库（不修改 status、tree_json）
        await session.execute(
            text(
                """
                UPDATE sop_document
                SET variable_schema = CAST(:variable_schema AS jsonb),
                    updated_at = :updated_at
                WHERE id = :id
                """
            ),
            {
                "id": document_id,
                "variable_schema": json.dumps(current_schema, ensure_ascii=False),
                "updated_at": datetime.now(UTC),
            },
        )
        resource_revision = None
        if row["status"] == "published":
            resource_revision = await _publish_sop_revision(session, document_id, get_current_trace_id())
        await session.commit()

    logger.info(
        event="sop_variable_schema_updated",
        document_id=document_id,
        updated_fields=updated_count,
    )

    return {
        "success": True,
        "document_id": document_id,
        "updated": updated_count,
        "variable_schema": current_schema,
        "resource_revision": resource_revision,
    }


# ─────────────────────────────────────────────────────────────────────────────
# KBD 条目内容编辑接口
# ─────────────────────────────────────────────────────────────────────────────


ReviewedImageSeq = Annotated[StrictInt, Field(ge=0)]


class KbdUpdateRequest(BaseModel):
    """KBD 条目内容编辑请求

    支持编辑标题、8 大章节字段、signals_json 和分类。
    编辑章节字段后，content_md 自动由章节重建（不含视觉描述）。
    若明确传入 content_md，则优先使用传入的值。
    """

    title: str | None = Field(None, max_length=500, description="新标题（可选）")
    # 8 大章节字段
    problem_description: str | None = Field(None, description="问题描述章节")
    alert_info: str | None = Field(None, description="告警信息章节")
    steps_text: str | None = Field(None, description="有效排查步骤（自然语言 Markdown）")
    root_cause: str | None = Field(None, description="根因章节")
    solution: str | None = Field(None, description="解决方案章节")
    operational_impact: str | None = Field(None, description="操作影响范围章节")
    is_temporary: str | None = Field(None, description="是否是临时解决方案章节")
    recommendations: str | None = Field(None, description="建议与总结章节")
    # 关键信号集合（agent 可执行与判定）
    # v2 列形态（RFC §7）：接受 v2 数组级对象 {schema_version, signals}。
    signals_json: Any | None = Field(
        None,
        description="关键信号集合：v2 对象 {schema_version,signals}",
    )
    delete_signal_id: str | None = Field(
        None,
        min_length=1,
        max_length=256,
        description="按稳定 Signal ID 删除权威工作稿中的单条信号；不得与 signals_json 同时提交",
    )
    images_json: list[dict[str, Any]] | None = Field(
        None,
        description="截图 Evidence；按稳定 seq 整体保存",
    )
    reviewed_image_seqs: list[ReviewedImageSeq] | None = Field(
        None,
        description=(
            "本次明确由专家确认的图片 seq；只提升这些图片为 expert_confirmed。"
            "兼容旧调用方：省略时表示提交的全部图片均已审核"
        ),
    )
    # 聚合渲染（可选，不传则自动由章节重建）
    content_md: str | None = Field(None, description="聚合 Markdown（优先用传入的值；不传则自动由章节重建）")
    content_raw: str | None = Field(None, description="新纯文本去噪内容（可选）")
    category_id: str | None = Field(None, description="新分类 ID（可选）")
    change_annotations: list[dict[str, Any]] | None = Field(
        None,
        description="专家修改原因标注；支持 path 或稳定 signal_id，供后续 LLM 评估使用",
    )
    lock_version: int | None = Field(None, ge=0, description="可选乐观锁版本；不匹配时返回 409")


def _mark_payload_signal_generation_stale(payload: dict[str, Any]) -> None:
    """知识来源变化时把工作稿 Signal 标为过期，但保留生成追溯元数据。"""

    signals_doc = _load_signals_json(payload.get("signals_json"))
    metadata = signals_doc.get("generation_metadata")
    if isinstance(metadata, dict):
        metadata["status"] = "stale"
        payload["signals_json"] = signals_doc


def _normalize_maintenance_images(
    raw_images: list[dict[str, Any]],
    reviewed_image_seqs: list[int] | None = None,
) -> list[dict[str, Any]]:
    """校验截图 Evidence，并只提升本次明确审核的图片。

    ``images_json`` 是整体保存契约，但一次 UI 操作只编辑一张图。新版调用方通过
    ``reviewed_image_seqs`` 声明本次真实审核范围，避免把未打开的其他图片连带标成
    ``expert_confirmed``。省略该字段时保留历史“整组均已审核”语义。
    """

    normalized: list[dict[str, Any]] = []
    seen_seqs: set[int] = set()
    reviewed_seqs = set(reviewed_image_seqs) if reviewed_image_seqs is not None else None
    if reviewed_seqs is not None and len(reviewed_seqs) != len(reviewed_image_seqs):
        raise HTTPException(status_code=422, detail="reviewed_image_seqs 不能包含重复 seq")
    for index, raw_image in enumerate(raw_images):
        if not isinstance(raw_image, dict):
            raise HTTPException(status_code=422, detail=f"images_json[{index}] 必须是对象")
        image = copy.deepcopy(raw_image)
        seq = image.get("seq")
        if not isinstance(seq, int) or isinstance(seq, bool) or seq < 0:
            raise HTTPException(status_code=422, detail=f"images_json[{index}].seq 必须是非负整数")
        if seq in seen_seqs:
            raise HTTPException(status_code=422, detail=f"images_json 存在重复 seq={seq}")
        seen_seqs.add(seq)
        if not isinstance(image.get("section"), str) or not image["section"].strip():
            raise HTTPException(status_code=422, detail=f"images_json[{index}].section 不能为空")
        if not isinstance(image.get("desc"), str):
            raise HTTPException(status_code=422, detail=f"images_json[{index}].desc 必须是字符串")
        evidence = image.get("evidence")
        if "evidence" in image and not isinstance(evidence, dict):
            raise HTTPException(status_code=422, detail=f"images_json[{index}].evidence 必须是对象")
        if isinstance(evidence, dict):
            if "quality" in evidence and not isinstance(evidence["quality"], dict):
                raise HTTPException(status_code=422, detail=f"images_json[{index}].evidence.quality 必须是对象")
            if "provenance" in evidence and not isinstance(evidence["provenance"], dict):
                raise HTTPException(status_code=422, detail=f"images_json[{index}].evidence.provenance 必须是对象")
        if reviewed_seqs is None or seq in reviewed_seqs:
            if evidence is None:
                evidence = {}
                image["evidence"] = evidence
            quality = evidence.setdefault("quality", {})
            quality.update(
                {
                    "status": "manual_reviewed",
                    "needs_review": False,
                    "inference_status": "expert_confirmed",
                    "inference_needs_review": False,
                }
            )
            provenance = evidence.setdefault("provenance", {})
            if not isinstance(provenance, dict):
                provenance = {}
                evidence["provenance"] = provenance
            provenance["expert_edited"] = True
        normalized.append(image)
    if reviewed_seqs is not None:
        unknown_seqs = reviewed_seqs - seen_seqs
        if unknown_seqs:
            unknown_text = ", ".join(str(seq) for seq in sorted(unknown_seqs))
            raise HTTPException(
                status_code=422,
                detail=f"reviewed_image_seqs 包含 images_json 中不存在的 seq: {unknown_text}",
            )
    return sorted(normalized, key=lambda item: item["seq"])


def _patch_maintenance_payload(payload: dict[str, Any], body: KbdUpdateRequest) -> dict[str, Any]:
    """把 Admin Patch 应用到独立维护 payload，不触碰已发布主记录。"""

    if body.signals_json is not None and body.delete_signal_id is not None:
        raise HTTPException(status_code=400, detail="signals_json 与 delete_signal_id 不能同时提交")
    result = copy.deepcopy(payload)
    section_fields = KbdEntry.SECTION_FIELDS
    source_changed = any(
        getattr(body, field) is not None
        for field in ("title", "problem_description", "alert_info", "steps_text", "category_id", "images_json")
    )
    for field in ("title", *section_fields, "content_raw", "category_id"):
        value = getattr(body, field)
        if value is not None:
            result[field] = value

    if body.signals_json is not None:
        result["signals_json"] = _prepare_expert_draft_signals(body.signals_json)
    elif body.delete_signal_id is not None:
        result["signals_json"] = _delete_signal_from_document(result.get("signals_json"), body.delete_signal_id)

    old_images = copy.deepcopy(result.get("images_json") or [])
    if body.images_json is not None:
        result["images_json"] = _normalize_maintenance_images(
            body.images_json,
            body.reviewed_image_seqs,
        )

    if body.content_md is not None:
        result["content_md"] = body.content_md
        result["content_raw"] = body.content_raw or strip_markdown(body.content_md)
    elif any(getattr(body, field) is not None for field in section_fields) or body.images_json is not None:
        working_entry = KbdEntry()
        for field in KBD_PAYLOAD_FIELDS:
            if field in result:
                setattr(working_entry, field, result[field])
        working_entry.entry_metadata = result.get("metadata") or {}
        result["content_md"] = working_entry.rebuild_content_md(old_images_json=old_images)
        result["content_raw"] = strip_markdown(result["content_md"] or "")

    if source_changed and body.signals_json is None:
        _mark_payload_signal_generation_stale(result)
    result["payload_schema_version"] = 1
    return result


def _working_payload_response(kbd: KbdEntry, revision: KbdRevision) -> dict[str, Any]:
    """返回前端可直接覆盖当前审核视图的工作稿内容。"""

    payload = revision.payload_json
    return {
        "success": True,
        "kbd_id": kbd.id,
        "status": kbd.status,
        "maintenance_working": True,
        "working_revision_id": revision.id,
        "lock_version": kbd.lock_version,
        "payload": payload,
        "knowledge_revision": revision_metadata(revision),
        "resource_revision": None,
        "agent_active_unchanged": True,
    }


@kbd_router.post("/{kbd_id}/maintenance")
async def create_kbd_maintenance_working(request: Request, kbd_id: int) -> dict[str, Any]:
    """为已发布 KBD 创建独立维护工作稿；不改变 Agent active。"""

    _check_auth(request)
    if _db_manager is None:
        raise HTTPException(status_code=503, detail="数据库未就绪")
    async with _db_manager.async_session_factory() as session:
        kbd = await session.get(KbdEntry, kbd_id, with_for_update=True)
        if kbd is None:
            raise HTTPException(status_code=404, detail=f"KBD 条目 {kbd_id} 不存在")
        if kbd.status != "published":
            raise HTTPException(status_code=409, detail="只有已发布 KBD 需要创建维护工作稿")
        if kbd.working_revision_id is not None:
            current = await session.get(KbdRevision, kbd.working_revision_id)
            if current is not None and (current.generation_metadata or {}).get("origin") == "admin_maintenance":
                return _working_payload_response(kbd, current)
        revision = await ensure_kbd_revision_payload(
            session,
            kbd=kbd,
            payload=build_kbd_revision_payload(kbd),
            revision_type="expert",
            actor_type="system",
            parent_revision_id=kbd.working_revision_id or kbd.latest_proposal_revision_id,
            generation_metadata={"origin": "admin_maintenance", "status": "opened"},
            validation_summary={"status": "not_run"},
            trace_id=get_current_trace_id(),
            reuse_existing=False,
        )
        kbd.lock_version += 1
        await session.commit()
        return _working_payload_response(kbd, revision)


@kbd_router.patch("/{kbd_id}/maintenance")
async def update_kbd_maintenance_working(
    request: Request,
    kbd_id: int,
    body: KbdUpdateRequest,
) -> dict[str, Any]:
    """保存已发布 KBD 的独立维护工作稿；不覆盖 active 主记录。"""

    _check_auth(request)
    if _db_manager is None:
        raise HTTPException(status_code=503, detail="数据库未就绪")
    async with _db_manager.async_session_factory() as session:
        kbd = await session.get(KbdEntry, kbd_id, with_for_update=True)
        if kbd is None:
            raise HTTPException(status_code=404, detail=f"KBD 条目 {kbd_id} 不存在")
        if kbd.status != "published" or kbd.working_revision_id is None:
            raise HTTPException(status_code=409, detail="请先创建维护工作稿")
        if body.lock_version is not None and body.lock_version != kbd.lock_version:
            raise HTTPException(
                status_code=409,
                detail={"code": "KBD_EDIT_CONFLICT", "message": "维护工作稿已被更新，请刷新后重试"},
            )
        current = await session.get(KbdRevision, kbd.working_revision_id)
        if current is None or (current.generation_metadata or {}).get("origin") != "admin_maintenance":
            raise HTTPException(status_code=409, detail="当前不存在可编辑的维护工作稿")
        try:
            annotations = normalize_change_annotations(body.change_annotations)
        except ValueError as exc:
            raise HTTPException(status_code=422, detail=str(exc)) from exc
        payload = _patch_maintenance_payload(current.payload_json, body)
        revision = await ensure_kbd_revision_payload(
            session,
            kbd=kbd,
            payload=payload,
            revision_type="expert",
            actor_type="expert",
            actor_id=None,
            parent_revision_id=current.id,
            generation_metadata={"origin": "admin_maintenance", "status": "saved", "identity_status": "unavailable"},
            validation_summary={"status": "saved", "publish_validation": "not_run"},
            review_metadata=build_review_metadata(
                parent_payload=current.payload_json,
                payload=payload,
                annotations=annotations,
                identity_status="unavailable",
                review_state="working",
            ),
            trace_id=get_current_trace_id(),
        )
        kbd.lock_version += 1
        await session.commit()
        return _working_payload_response(kbd, revision)


@kbd_router.delete("/{kbd_id}/maintenance")
async def discard_kbd_maintenance_working(request: Request, kbd_id: int) -> dict[str, Any]:
    """放弃未发布维护稿；append-only 历史保留，Agent active 不变。"""

    _check_auth(request)
    if _db_manager is None:
        raise HTTPException(status_code=503, detail="数据库未就绪")
    async with _db_manager.async_session_factory() as session:
        kbd = await session.get(KbdEntry, kbd_id, with_for_update=True)
        if kbd is None:
            raise HTTPException(status_code=404, detail=f"KBD 条目 {kbd_id} 不存在")
        kbd.working_revision_id = None
        kbd.lock_version += 1
        await session.commit()
    return {"success": True, "kbd_id": kbd_id, "agent_active_unchanged": True}


@kbd_router.post("/{kbd_id}/maintenance/publish", response_model=KbdApproveResponse)
async def publish_kbd_maintenance_working(
    request: Request,
    kbd_id: int,
    body: KbdApproveRequest,
) -> KbdApproveResponse:
    """发布维护工作稿并原子切换 Agent active；失败时旧 active 保持不变。"""

    _check_auth(request)
    if _db_manager is None:
        raise HTTPException(status_code=503, detail="数据库未就绪")

    async with _db_manager.async_session_factory() as session:
        kbd = await session.get(KbdEntry, kbd_id)
        if kbd is None:
            raise HTTPException(status_code=404, detail=f"KBD 条目 {kbd_id} 不存在")
        if kbd.status != "published" or kbd.working_revision_id is None:
            raise HTTPException(status_code=409, detail="当前没有待发布的维护工作稿")
        working_revision_id = kbd.working_revision_id
        source_lock_version = kbd.lock_version
        working = await session.get(KbdRevision, working_revision_id)
        if working is None or (working.generation_metadata or {}).get("origin") != "admin_maintenance":
            raise HTTPException(status_code=409, detail="当前工作版本不是维护工作稿")
        payload = copy.deepcopy(working.payload_json)

    if body.lock_version is not None and body.lock_version != source_lock_version:
        raise HTTPException(
            status_code=409,
            detail={"code": "KBD_EDIT_CONFLICT", "message": "维护工作稿已被更新，请刷新后重新发布"},
        )
    if body.category_id:
        payload["category_id"] = body.category_id
    for field, label in (
        ("title", "标题"),
        ("problem_description", "问题描述"),
        ("root_cause", "根因"),
        ("solution", "解决方案"),
    ):
        if not str(payload.get(field) or "").strip():
            raise HTTPException(status_code=422, detail=f"{label}不能为空")
    signals_doc = _load_signals_json(payload.get("signals_json"))
    try:
        signals_doc = _prepare_expert_publish_signals(signals_doc)
    except jsonschema.ValidationError as exc:
        _raise_signal_validation_error(exc, signals_doc.get("signals") or [])
    payload["signals_json"] = signals_doc
    signals = signals_doc.get("signals") or []
    if not any(
        isinstance(signal, dict)
        and (
            str((signal.get("acquire") or {}).get("tool") or "").startswith("qfk")
            or (signal.get("provenance") or {}).get("category") == "backend"
        )
        for signal in signals
    ):
        raise HTTPException(status_code=422, detail="至少需要一条消费者 QFK 信号")
    effective_category_id = payload.get("category_id") or payload.get("ai_category_id")
    if not effective_category_id:
        raise HTTPException(status_code=422, detail="请先确认分类")
    payload["category_id"] = effective_category_id
    content_md = str(payload.get("content_md") or "")
    if not content_md:
        raise HTTPException(status_code=422, detail="维护工作稿缺少完整内容")
    embedding_text = build_kbd_embedding_text(
        title=payload.get("title"),
        problem_description=payload.get("problem_description"),
        alert_info=payload.get("alert_info"),
        root_cause=payload.get("root_cause"),
        fallback_text=payload.get("content_raw") or content_md,
    )
    payload["content_raw"] = strip_markdown(content_md)

    embedding_vector: list[float] | None = None
    if _embedding_service:
        try:
            embedding_vector = await _embedding_service.embed_single(embedding_text)
        except Exception as exc:
            logger.warning(event="kbd_maintenance_embedding_failed", kbd_id=kbd_id, error=str(exc))
    embedding_content_hash = hashlib.sha256(embedding_text.encode("utf-8")).hexdigest()
    now = datetime.now(UTC)
    tsv_text = segment(f"{payload.get('title', '')} {payload['content_raw']}")

    async with _db_manager.async_session_factory() as session:
        kbd = await session.get(KbdEntry, kbd_id, with_for_update=True)
        if (
            kbd is None
            or kbd.status != "published"
            or kbd.working_revision_id != working_revision_id
            or kbd.lock_version != source_lock_version
        ):
            raise HTTPException(status_code=409, detail="发布期间维护工作稿已变化，旧生效版未受影响")
        working = await session.get(KbdRevision, working_revision_id)
        if working is None:
            raise HTTPException(status_code=409, detail="维护工作稿不存在")
        apply_kbd_revision_payload(kbd, payload)
        kbd.reviewer_id = body.reviewer_id
        kbd.reviewed_at = now
        kbd.review_note = body.review_note or kbd.review_note
        kbd.published_at = now
        kbd.lock_version += 1
        await session.flush()

        embedding_sql = (
            "embedding = CAST(:embedding AS vector), embedding_model = :embedding_model, "
            "embedding_content_hash = :embedding_content_hash, embedding_updated_at = :embedding_updated_at"
            if embedding_vector
            else "embedding = NULL, embedding_model = NULL, embedding_content_hash = NULL, embedding_updated_at = NULL"
        )
        params: dict[str, Any] = {"id": kbd_id, "tsv_text": tsv_text}
        if embedding_vector:
            params.update(
                {
                    "embedding": "[" + ",".join(str(value) for value in embedding_vector) + "]",
                    "embedding_model": _embedding_service.model_name,
                    "embedding_content_hash": embedding_content_hash,
                    "embedding_updated_at": now,
                }
            )
        await session.execute(
            text(f"UPDATE kbd_entry SET {embedding_sql}, tsv = to_tsvector('simple', :tsv_text) WHERE id = :id"),  # noqa: S608
            params,
        )
        approved_revision = await ensure_kbd_revision_payload(
            session,
            kbd=kbd,
            payload=build_kbd_revision_payload(kbd),
            revision_type="expert",
            actor_type="expert",
            actor_id=None,
            parent_revision_id=working_revision_id,
            generation_metadata={
                "origin": "admin_maintenance_publish",
                "identity_status": "unverified_body_reviewer_id",
                "legacy_reviewer_id": body.reviewer_id,
                "review_note": body.review_note or "",
            },
            validation_summary={"status": "passed", "gate": "publishable_signals_json"},
            review_metadata=build_review_metadata(
                parent_payload=working.payload_json,
                payload=build_kbd_revision_payload(kbd),
                identity_status="unverified_body_reviewer_id",
                review_state="approved",
                reviewer_id=body.reviewer_id,
                review_note=body.review_note,
            ),
            trace_id=get_current_trace_id(),
            reuse_existing=False,
        )
        resource_revision = await _publish_kbd_revision(session, kbd_id, get_current_trace_id())
        kbd.working_revision_id = None
        await session.commit()

    return KbdApproveResponse(
        success=True,
        kbd_id=kbd_id,
        status="published",
        embedding_generated=embedding_vector is not None,
        published_at=now.isoformat(),
        resource_revision=resource_revision,
        knowledge_revision=revision_metadata(approved_revision),
    )


@kbd_router.patch("/{kbd_id}")
async def update_kbd_entry(request: Request, kbd_id: int, body: KbdUpdateRequest):
    """编辑 KBD 条目的标题、章节字段、signals_json 或分类。

    处理逻辑：
    1. 章节字段任意一个被修改时，如果没有明确提供 content_md，则自动先从数据库读取当前章节状态并重建 content_md
    2. 如果明确提供了 content_md，则用传入的值（优先）
    """
    _check_auth(request)

    if _db_manager is None:
        raise HTTPException(status_code=503, detail="数据库未就绪")
    if body.signals_json is not None and body.delete_signal_id is not None:
        raise HTTPException(status_code=400, detail="signals_json 与 delete_signal_id 不能同时提交")

    # kbd_entry 仍是 Agent 当前生效内容的兼容主记录。已发布维护必须走独立 maintenance
    # working；普通 PATCH 不得把未审核内容写入主记录或静默切换 active。
    async with _db_manager.async_session_factory() as session:
        current_result = await session.execute(select(KbdEntry.status).where(KbdEntry.id == kbd_id))
        current_status = current_result.scalar_one_or_none()
        if current_status is None:
            raise HTTPException(status_code=404, detail=f"KBD 条目 {kbd_id} 不存在")
        if current_status == "published":
            raise HTTPException(
                status_code=409,
                detail="已发布 KBD 不能直接覆盖编辑；请先创建维护工作稿，当前 Agent 生效版保持不变",
            )

    # 所有可更新字段
    section_fields = (
        "problem_description",
        "alert_info",
        "steps_text",
        "root_cause",
        "solution",
        "operational_impact",
        "is_temporary",
        "recommendations",
    )
    any_section_changed = any(getattr(body, f) is not None for f in section_fields)
    signal_source_changed = (
        body.title is not None
        or body.problem_description is not None
        or body.alert_info is not None
        or body.steps_text is not None
        or body.category_id is not None
        or body.images_json is not None
    )
    has_any_field = (
        body.title is not None
        or any_section_changed
        or body.signals_json is not None
        or body.delete_signal_id is not None
        or body.images_json is not None
        or body.content_md is not None
        or body.content_raw is not None
        or body.category_id is not None
    )
    if not has_any_field:
        raise HTTPException(status_code=400, detail="至少需要提供一个可更新字段")

    set_clauses = []
    params: dict = {"id": kbd_id}

    if body.title is not None:
        set_clauses.append("title = :title")
        params["title"] = body.title

    for field in section_fields:
        val = getattr(body, field)
        if val is not None:
            set_clauses.append(f"{field} = :{field}")
            params[field] = val

    if body.signals_json is not None:
        # 直接切 v2 列形态（RFC §7）：保存时统一归约为 v2 数组级对象
        v2_doc = _prepare_expert_draft_signals(body.signals_json)
        # 必须用 CAST(:signals_json AS jsonb)，不能写成 ":signals_json::jsonb"。
        # 后者中 ':signals_json' 紧跟 '::'，SQLAlchemy 命名绑定正则(负向预查 (?!:))
        # 不把它识别为绑定参数，会原样发给 Postgres 触发 'syntax error at or near ":"' (500)，
        # 前端统一弹「保存失败，请重试」——即 PR#599 修过、PR#601 回退复现的坑。
        set_clauses.append("signals_json = CAST(:signals_json AS jsonb)")
        params["signals_json"] = json.dumps(v2_doc, ensure_ascii=False)

    normalized_images: list[dict[str, Any]] | None = None
    if body.images_json is not None:
        normalized_images = _normalize_maintenance_images(
            body.images_json,
            body.reviewed_image_seqs,
        )
        set_clauses.append("images_json = CAST(:images_json AS jsonb)")
        params["images_json"] = json.dumps(normalized_images, ensure_ascii=False)
    elif signal_source_changed:
        # 章节/分类一旦变化，旧 Signal/Contract 的来源指纹立即失效；保留原数据供
        # 审核 diff，但运行时 Compiler 会 fail closed，直到重新抽取或人工审核。
        set_clauses.append(
            "signals_json = CASE "
            "WHEN signals_json ? 'generation_metadata' "
            "THEN jsonb_set(signals_json, '{generation_metadata,status}', '\"stale\"'::jsonb, true) "
            "ELSE signals_json END"
        )

    # content_md 处理：明确传入则用传入的值；有章节更改则需先读库并重建
    if body.content_md is not None:
        # 明确传入了 content_md，优先使用
        # 为了保障 8 大章节字段与 content_md 保持 100% 一致，读取库中 images_json 并进行解析同步
        async with _db_manager.async_session_factory() as session:
            cur_result = await session.execute(
                text("SELECT images_json FROM kbd_entry WHERE id = :id"),
                {"id": kbd_id},
            )
            cur_row = cur_result.mappings().first()
            if not cur_row:
                raise HTTPException(status_code=404, detail=f"KBD 条目 {kbd_id} 不存在")
            images_json = cur_row["images_json"] or []

        # 实例化临时 KbdEntry，利用已实现的 sync_sections_from_content_md 进行反向解析
        temp_entry = KbdEntry(
            content_md=body.content_md,
            images_json=images_json,
        )
        temp_entry.sync_sections_from_content_md()

        # 将解析出的 8 大章节字段也一并放入 SQL 更新中
        for field in section_fields:
            parsed_val = getattr(temp_entry, field)
            if f"{field} = :{field}" not in set_clauses:
                set_clauses.append(f"{field} = :{field}")
            params[field] = parsed_val

        if "content_md = :content_md" not in set_clauses:
            set_clauses.append("content_md = :content_md")
        params["content_md"] = body.content_md
        if "content_raw = :content_raw" not in set_clauses:
            set_clauses.append("content_raw = :content_raw")
        params["content_raw"] = body.content_raw or strip_markdown(body.content_md)
    elif any_section_changed:
        # 章节有变更且未传入 content_md：读库 + 应用 patch + 重建
        async with _db_manager.async_session_factory() as session:
            cur_result = await session.execute(
                text(
                    "SELECT problem_description, alert_info, steps_text, root_cause, "
                    "solution, operational_impact, is_temporary, recommendations "
                    "FROM kbd_entry WHERE id = :id"
                ),
                {"id": kbd_id},
            )
            cur_row = cur_result.mappings().first()
            if not cur_row:
                raise HTTPException(status_code=404, detail=f"KBD 条目 {kbd_id} 不存在")

        # 合并：优先用 body 中的新值，其次用库中现有值
        section_map = {
            "problem_description": "问题描述",
            "alert_info": "告警信息",
            "steps_text": "有效排查步骤",
            "root_cause": "根因",
            "solution": "解决方案",
            "operational_impact": "操作影响范围",
            "is_temporary": "是否是临时解决方案",
            "recommendations": "建议与总结",
        }
        merged_parts = []
        for field, heading in section_map.items():
            # body 中有新值用新值，否则用库中现有值
            text_val = (getattr(body, field) if getattr(body, field) is not None else cur_row[field] or "").strip()
            if text_val:
                merged_parts.append(f"## {heading}\n\n{text_val}")
        rebuilt_content_md = "\n\n".join(merged_parts)
        set_clauses.append("content_md = :content_md")
        params["content_md"] = rebuilt_content_md
        set_clauses.append("content_raw = :content_raw")
        params["content_raw"] = body.content_raw or strip_markdown(rebuilt_content_md)
    elif body.content_raw is not None:
        set_clauses.append("content_raw = :content_raw")
        params["content_raw"] = body.content_raw

    if body.category_id is not None:
        set_clauses.append("category_id = :category_id")
        params["category_id"] = body.category_id

    async with _db_manager.async_session_factory() as session:
        kbd_result = await session.execute(select(KbdEntry).where(KbdEntry.id == kbd_id).with_for_update())
        kbd = kbd_result.scalar_one_or_none()
        if kbd is None:
            raise HTTPException(status_code=404, detail=f"KBD 条目 {kbd_id} 不存在")
        if kbd.status == "published":
            raise HTTPException(
                status_code=409,
                detail="已发布 KBD 不能直接覆盖编辑；请等待 working revision 维护入口，当前 active 保持不变",
            )
        if body.lock_version is not None and body.lock_version != kbd.lock_version:
            raise HTTPException(
                status_code=409,
                detail={
                    "code": "KBD_EDIT_CONFLICT",
                    "message": "该 KBD 已被其他编辑更新，请刷新后重试",
                    "expected_lock_version": body.lock_version,
                    "current_lock_version": kbd.lock_version,
                },
            )
        try:
            annotations = normalize_change_annotations(body.change_annotations)
        except ValueError as exc:
            raise HTTPException(status_code=422, detail=str(exc)) from exc
        old_images_json = list(kbd.images_json or [])

        if body.delete_signal_id is not None:
            v2_doc = _delete_signal_from_document(kbd.signals_json, body.delete_signal_id)
            set_clauses.append("signals_json = CAST(:signals_json AS jsonb)")
            params["signals_json"] = json.dumps(v2_doc, ensure_ascii=False)

        if kbd.latest_proposal_revision_id is None:
            signal_metadata = _load_signals_json(kbd.signals_json).get("generation_metadata") or {}
            proposal_revision = await ensure_kbd_revision(
                session,
                kbd=kbd,
                revision_type="proposal",
                actor_type="llm" if signal_metadata.get("model_id") else "migration",
                generation_metadata={"origin": "pre_expert_edit_baseline", **signal_metadata},
                trace_id=get_current_trace_id(),
            )
        else:
            proposal_revision = await session.get(KbdRevision, kbd.latest_proposal_revision_id)
        parent_revision_id = kbd.working_revision_id or (
            proposal_revision.id if proposal_revision is not None else None
        )
        # 工作稿允许连续保存。训练/评估数据中的字段差异必须相对“这一次编辑的
        # 直接父版本”计算，不能一直相对最初 Proposal；否则第二次保存会把此前已经
        # 标注过的改动重复记为本次专家修正，reason_code 将失真。
        parent_revision = await session.get(KbdRevision, parent_revision_id) if parent_revision_id else None

        set_clauses.append("lock_version = lock_version + 1")
        result = await session.execute(
            text(f"UPDATE kbd_entry SET {', '.join(set_clauses)} WHERE id = :id RETURNING id, status, lock_version"),  # noqa: S608
            params,
        )
        updated = result.mappings().first()
        if not updated:
            raise HTTPException(status_code=404, detail=f"KBD 条目 {kbd_id} 不存在")
        await session.refresh(kbd)
        if normalized_images is not None:
            # 章节字段仍保存 ![img:N] 占位符；使用旧 Evidence 还原历史展开块，再以专家
            # 确认后的 Evidence 重建 Agent 可见 content_md。
            kbd.content_md = kbd.rebuild_content_md(old_images_json=old_images_json)
            kbd.content_raw = strip_markdown(kbd.content_md or "")
            await session.flush()
        expert_revision = await ensure_kbd_revision(
            session,
            kbd=kbd,
            revision_type="expert",
            actor_type="expert",
            actor_id=None,
            parent_revision_id=parent_revision_id,
            generation_metadata={"origin": "admin_working_edit", "identity_status": "unavailable"},
            validation_summary={"status": "saved", "publish_validation": "not_run"},
            review_metadata=build_review_metadata(
                parent_payload=getattr(parent_revision, "payload_json", None) if parent_revision is not None else {},
                payload=build_kbd_revision_payload(kbd),
                annotations=annotations,
                identity_status="unavailable",
                review_state="working",
            ),
            trace_id=get_current_trace_id(),
        )
        await session.commit()

    logger.info(event="kbd_updated", kbd_id=kbd_id, fields=list(params.keys()))
    return {
        "success": True,
        "kbd_id": kbd_id,
        "signals_json": _signals_for_response(kbd.signals_json),
        "images_json": kbd.images_json or [],
        "content_md": kbd.content_md or "",
        "lock_version": updated.get("lock_version", getattr(kbd, "lock_version", 0)),
        "knowledge_revision": revision_metadata(expert_revision),
        "resource_revision": None,
    }


# ─────────────────────────────────────────────────────────────────────────────
# KBD 条目重新发布接口（rejected → published）
# ─────────────────────────────────────────────────────────────────────────────


@kbd_router.post("/{kbd_id}/republish", response_model=KbdApproveResponse)
async def republish_kbd_entry(request: Request, kbd_id: int, body: KbdApproveRequest):
    """重新发布已拒绝的 KBD 条目（rejected → published），重新生成 embedding"""
    _check_auth(request)

    if _db_manager is None:
        raise HTTPException(status_code=503, detail="数据库未就绪")

    logger.info(event="kbd_republish_request", kbd_id=kbd_id, reviewer_id=body.reviewer_id)

    # 查询条目（允许 rejected 或 draft 状态）
    async with _db_manager.async_session_factory() as session:
        result = await session.execute(
            text(
                "SELECT id, title, content_md, content_raw, problem_description, alert_info, root_cause, status, "
                "category_id, ai_category_id, signals_json, lock_version FROM kbd_entry WHERE id = :id"
            ),
            {"id": kbd_id},
        )
        row = result.mappings().first()
        if not row:
            raise HTTPException(status_code=404, detail=f"KBD 条目 {kbd_id} 不存在")
        if row["status"] not in {"draft", "rejected"}:
            raise HTTPException(
                status_code=400,
                detail=f"KBD 条目当前状态为 {row['status']}，只有 draft/rejected 状态可重新发布",
            )
        source_lock_version = int(row.get("lock_version") or 0)
        if body.lock_version is not None and body.lock_version != source_lock_version:
            raise HTTPException(
                status_code=409,
                detail={
                    "code": "KBD_EDIT_CONFLICT",
                    "message": "该 KBD 已被其他编辑更新，请刷新后重新审核",
                    "expected_lock_version": body.lock_version,
                    "current_lock_version": source_lock_version,
                },
            )
        content_md = row["content_md"]
        if not content_md:
            raise HTTPException(status_code=400, detail=f"KBD 条目 {kbd_id} 缺少 content_md")
        _republish_doc = _load_signals_json(row["signals_json"]) or {}
        _republish_signals = _republish_doc.get("signals", [])
        if not _republish_signals:
            raise HTTPException(
                status_code=422,
                detail=f"KBD 条目 {kbd_id} 缺少关键信号，禁止重新发布",
            )
        try:
            _republish_doc = _prepare_expert_publish_signals(_republish_doc)
        except jsonschema.ValidationError as exc:
            _raise_signal_validation_error(exc, _republish_doc.get("signals") or [])
        embedding_text = build_kbd_embedding_text(
            title=row["title"],
            problem_description=row["problem_description"],
            alert_info=row["alert_info"],
            root_cause=row["root_cause"],
            fallback_text=row["content_raw"] or content_md,
        )

    # 生成 embedding（事务外调用）
    embedding_generated = False
    embedding_vector: list[float] | None = None
    embedding_content_hash = hashlib.sha256(embedding_text.encode("utf-8")).hexdigest()
    if _embedding_service:
        try:
            embedding_vector = await _embedding_service.embed_single(embedding_text)
            embedding_generated = True
            logger.info(event="kbd_republish_embedding_generated", kbd_id=kbd_id, vector_dim=len(embedding_vector))
        except Exception as exc:
            logger.warning(event="kbd_republish_embedding_failed", kbd_id=kbd_id, error=str(exc))

    now = datetime.now(UTC)
    current_content_raw = strip_markdown(content_md or "")
    tsv_text = segment(f"{row['title']} {current_content_raw}")
    # 门 2（重发布同样适用）：category_id 与 ai_category_id 同步，根治孤儿 KBD
    # 优先采用人工确认的分类（发布请求 body.category_id），fallback 到 DB 已有值
    effective_category_id = body.category_id or row["category_id"] or row["ai_category_id"]
    if not effective_category_id:
        raise HTTPException(
            status_code=422,
            detail=f"KBD 条目 {kbd_id} 缺少分类（category_id 与 ai_category_id 均为空），"
            f"请先抽取分类或在发布时确认分类后再重新发布",
        )
    async with _db_manager.async_session_factory() as session:
        if embedding_vector:
            vector_str = "[" + ",".join(str(v) for v in embedding_vector) + "]"
            update_sql = text(
                """
                UPDATE kbd_entry
                SET status = 'published',
                    published_at = :published_at,
                    reviewer_id = :reviewer_id,
                    reviewed_at = :reviewed_at,
                    review_note = COALESCE(:review_note, review_note),
                    content_raw = :content_raw,
                    category_id = :category_id,
                    signals_json = CAST(:signals_json AS jsonb),
                    embedding = CAST(:embedding AS vector),
                    embedding_model = :embedding_model,
                    embedding_content_hash = :embedding_content_hash,
                    embedding_updated_at = :embedding_updated_at,
                    tsv = to_tsvector('simple', :tsv_text)
                WHERE id = :id AND lock_version = :expected_lock_version AND status IN ('draft', 'rejected')
                RETURNING id, status, embedding, published_at
                """
            )
            params = {
                "id": kbd_id,
                "expected_lock_version": source_lock_version,
                "published_at": now,
                "reviewer_id": body.reviewer_id,
                "reviewed_at": now,
                "review_note": body.review_note,
                "content_raw": current_content_raw,
                "category_id": effective_category_id,
                "signals_json": json.dumps(_republish_doc, ensure_ascii=False),
                "embedding": vector_str,
                "embedding_model": _embedding_service.model_name,
                "embedding_content_hash": embedding_content_hash,
                "embedding_updated_at": now,
                "tsv_text": tsv_text,
            }
        else:
            update_sql = text(
                """
                UPDATE kbd_entry
                SET status = 'published',
                    published_at = :published_at,
                    reviewer_id = :reviewer_id,
                    reviewed_at = :reviewed_at,
                    review_note = COALESCE(:review_note, review_note),
                    content_raw = :content_raw,
                    category_id = :category_id,
                    signals_json = CAST(:signals_json AS jsonb),
                    embedding = NULL,
                    embedding_model = NULL,
                    embedding_content_hash = NULL,
                    embedding_updated_at = NULL,
                    tsv = to_tsvector('simple', :tsv_text)
                WHERE id = :id AND lock_version = :expected_lock_version AND status IN ('draft', 'rejected')
                RETURNING id, status, embedding, published_at
                """
            )
            params = {
                "id": kbd_id,
                "expected_lock_version": source_lock_version,
                "published_at": now,
                "reviewer_id": body.reviewer_id,
                "reviewed_at": now,
                "review_note": body.review_note,
                "content_raw": current_content_raw,
                "category_id": effective_category_id,
                "signals_json": json.dumps(_republish_doc, ensure_ascii=False),
                "tsv_text": tsv_text,
            }

        result = await session.execute(update_sql, params)
        updated = result.mappings().first()
        if not updated:
            raise HTTPException(
                status_code=409,
                detail={
                    "code": "KBD_EDIT_CONFLICT",
                    "message": "生成发布数据期间 KBD 已发生变化，未发布旧版本；请刷新后重新审核",
                    "expected_lock_version": source_lock_version,
                },
            )
        expert_revision = await _freeze_approved_expert_revision(
            session,
            kbd_id=kbd_id,
            reviewer_id=body.reviewer_id,
            review_note=body.review_note,
            trace_id=get_current_trace_id(),
        )
        resource_revision = await _publish_kbd_revision(session, kbd_id, get_current_trace_id())
        published_entry = await session.get(KbdEntry, kbd_id)
        if published_entry is not None:
            published_entry.working_revision_id = None
        await session.commit()

    logger.info(event="kbd_republished", kbd_id=kbd_id, reviewer_id=body.reviewer_id)
    return KbdApproveResponse(
        success=True,
        kbd_id=kbd_id,
        status="published",
        embedding_generated=embedding_generated,
        published_at=updated["published_at"].isoformat() if updated["published_at"] else None,
        resource_revision=resource_revision,
        knowledge_revision=revision_metadata(expert_revision),
    )


@kbd_router.post("/{kbd_id}/revert-to-draft")
async def revert_kbd_to_draft(request: Request, kbd_id: int):
    """将已发布/已拒绝的 KBD 条目退回待审核状态"""
    _check_auth(request)
    if _db_manager is None:
        raise HTTPException(status_code=503, detail="数据库未就绪")

    logger.info(event="kbd_revert_to_draft_request", kbd_id=kbd_id)
    async with _db_manager.async_session_factory() as session:
        result = await session.execute(text("SELECT id, status FROM kbd_entry WHERE id = :id"), {"id": kbd_id})
        row = result.mappings().first()
        if not row:
            raise HTTPException(status_code=404, detail=f"KBD 条目 {kbd_id} 不存在")
        if row["status"] == "draft":
            raise HTTPException(status_code=400, detail="当前已是待审核状态，无需操作")

        await session.execute(
            text("UPDATE kbd_entry SET status = 'draft', updated_at = NOW() WHERE id = :id"),
            {"id": kbd_id},
        )
        deactivated = await session.execute(
            text(
                "DELETE FROM dynamic_resource_active "
                "WHERE resource_type = 'kbd' AND resource_name = :resource_name"
            ),
            {"resource_name": str(kbd_id)},
        )
        await session.commit()

    logger.info(event="kbd_reverted_to_draft", kbd_id=kbd_id)
    return {
        "success": True,
        "kbd_id": kbd_id,
        "status": "draft",
        "active_deactivated": bool(deactivated.rowcount),
    }


# ─────────────────────────────────────────────────────────────────────────────
# KBD 在线重算 API（Prompt 修改后立即验证效果）
# ─────────────────────────────────────────────────────────────────────────────


@kbd_router.post("/{kbd_id}/reclassify", summary="重新分类单个 KBD 条目")
async def reclassify_kbd_entry(request: Request, kbd_id: int):
    """从 DB 读取 title + problem_desc，用最新 Prompt 重新分类。

    场景：admin-ui 修改 kbd_classify_v1 Prompt 后，点击"重新分类"按钮立即验证效果。
    更新字段：ai_category_id、ai_category_conf、ai_category_reason。
    """
    _check_auth(request)

    if _db_manager is None:
        raise HTTPException(status_code=503, detail="数据库未就绪")

    trace_id = get_current_trace_id()
    logger.info(event="kbd_reclassify_request", kbd_id=kbd_id, trace_id=trace_id)

    # 1. 读取 KBD 条目的 title 和 problem_desc
    async with _db_manager.async_session_factory() as session:
        entry = await _require_directly_mutable_kbd(session, kbd_id)
        title = entry.title or ""
        problem_desc = entry.problem_description or ""

    if not title:
        raise HTTPException(status_code=400, detail="KBD 条目缺少标题，无法分类")

    # 2. 调用分类核心逻辑（复用 classify.py 的 classify_case）
    from app.routes.classify import classify_case

    try:
        response = await classify_case(_db_manager, title, problem_desc)
    except HTTPException:
        raise
    except Exception as exc:
        logger.error(event="kbd_reclassify_failed", kbd_id=kbd_id, error=str(exc), trace_id=trace_id)
        raise HTTPException(status_code=500, detail=f"分类失败：{exc}")

    # 3. 更新 kbd_entry 的 AI 分类字段
    async with _db_manager.async_session_factory() as session:
        entry = await _require_directly_mutable_kbd(session, kbd_id, for_update=True)
        entry.ai_category_id = response.category_id
        entry.ai_category_conf = response.confidence
        entry.ai_category_reason = response.reason
        await session.commit()

    logger.info(
        event="kbd_reclassified",
        kbd_id=kbd_id,
        category_id=response.category_id,
        confidence=response.confidence,
        trace_id=trace_id,
    )

    return {
        "success": True,
        "kbd_id": kbd_id,
        "category_id": response.category_id,
        "confidence": response.confidence,
        "reason": response.reason,
        "needs_review": response.needs_review,
        "top3": [item.model_dump() for item in response.top3],
    }


@kbd_router.post("/{kbd_id}/reanalyze-images", summary="重新识图单个 KBD 条目（异步提交）")
async def reanalyze_kbd_images(request: Request, kbd_id: int):
    """异步提交 Vision 重算任务（Asynchronous Request-Reply 模式，P1-1）。

    场景：admin-ui 修改 kbd_vision_v1 Prompt 后，立即触发"重新识图"。
    更新字段：images_json、content_md（重建）。

    新行为：
      - 立即返回 202 + job_id（避免 HTTP 长连接阻塞/超时）
      - 后台 asyncio.create_task 执行实际 Vision LLM 调用
      - 客户端通过 GET /reanalyze-images/status?job_id=xxx 轮询状态

    旧同步行为（向后兼容 ?sync=true）：
      - sync=true 时沿用旧路径，直接返回完整结果（供 admin-ui 单条快速调试）

    Returns（异步）:
        {"job_id": str, "status": "pending", "kbd_id": int, "message": "已提交..."}
    """
    _check_auth(request)

    if _db_manager is None:
        raise HTTPException(status_code=503, detail="数据库未就绪")

    async with _db_manager.async_session_factory() as session:
        await _require_directly_mutable_kbd(session, kbd_id)

    trace_id = get_current_trace_id()

    # 向后兼容：?sync=true 直接同步返回（供 debug 场景）
    sync_mode = request.query_params.get("sync", "").lower() == "true"
    if sync_mode:
        return await _reanalyze_kbd_images_sync(kbd_id, trace_id)

    # 异步：提交 Job
    from app.services.vision_job_manager import get_job_manager
    from app.services.vision_processor import reanalyze_kbd_images as do_reanalyze

    async def _runner(k_id: int) -> dict[str, Any]:
        # 实时回写进度：避免轮询全程看到 running + 0/0 的「黑盒」观感；
        # 把本次请求的 trace_id 透传给 vision_processor，确保服务端日志与调用方串联。
        def _on_progress(done: int, failed: int, total: int) -> None:
            asyncio.create_task(jm._update(job_id, total=total, done=done, failed=failed))

        return await do_reanalyze(
            k_id,
            _db_manager.async_session_factory,
            on_progress=_on_progress,
            trace_id=trace_id,
        )

    jm = get_job_manager()
    job_id = await jm.submit(kbd_id, _runner, trace_id=trace_id)

    logger.info(
        event="kbd_reanalyze_images_submitted",
        kbd_id=kbd_id,
        job_id=job_id,
        trace_id=trace_id,
    )

    return JSONResponse(
        status_code=status.HTTP_202_ACCEPTED,
        content={
            "success": True,
            "kbd_id": kbd_id,
            "job_id": job_id,
            "status": "pending",
            "message": "Vision 任务已提交，请通过 GET /reanalyze-images/status 轮询进度",
        },
    )


async def _reanalyze_kbd_images_sync(kbd_id: int, trace_id: str):
    """旧同步路径（?sync=true 走此分支），保留兼容性。"""
    logger.info(event="kbd_reanalyze_images_sync_request", kbd_id=kbd_id, trace_id=trace_id)
    from app.services.vision_processor import reanalyze_kbd_images as do_reanalyze

    try:
        result = await do_reanalyze(kbd_id, _db_manager.async_session_factory)
    except PublishedKbdMutationError as exc:
        raise HTTPException(
            status_code=409,
            detail={"code": "KBD_MAINTENANCE_WORKING_REQUIRED", "message": str(exc)},
        ) from exc
    except ValueError as exc:
        raise HTTPException(status_code=404, detail=str(exc))
    except RuntimeError as exc:
        raise HTTPException(status_code=503, detail=str(exc))
    except Exception as exc:
        logger.error(event="kbd_reanalyze_images_failed", kbd_id=kbd_id, error=str(exc), trace_id=trace_id)
        raise HTTPException(status_code=500, detail=f"识图失败：{exc}")

    logger.info(
        event="kbd_reanalyze_images_completed",
        kbd_id=kbd_id,
        total=result["total"],
        done=result["done"],
        failed=result["failed"],
        trace_id=trace_id,
    )
    return {
        # 修复：同步路径不应恒为 True；部分图片失败时需如实反映 result.success
        "success": result.get("success", True),
        "kbd_id": kbd_id,
        "total": result["total"],
        "done": result["done"],
        "failed": result["failed"],
        "error": result.get("error"),
        "message": result.get("message", "识图完成"),
    }


@kbd_router.get("/{kbd_id}/reanalyze-images/status", summary="查询异步识图任务状态")
async def get_reanalyze_status(request: Request, kbd_id: int, job_id: str):
    """查询异步 Vision Job 状态。

    Args:
        kbd_id: KBD 条目 ID（路径参数，校验一致性）
        job_id: 任务 ID（来自 POST reanalyze-images 响应）

    Returns:
        {
            "job_id": str,
            "kbd_id": int,
            "status": "pending" | "running" | "done" | "failed",
            "total": int,
            "done": int,
            "failed": int,
            "error": str | None,
            "started_at": float | None,
            "finished_at": float | None,
        }
    """
    _check_auth(request)
    from app.services.vision_job_manager import get_job_manager

    job = await get_job_manager().get_status(job_id)
    if job is None:
        raise HTTPException(status_code=404, detail=f"Job {job_id} 不存在（可能已重启丢失）")
    if job["kbd_id"] != kbd_id:
        raise HTTPException(status_code=400, detail="job_id 与 kbd_id 不匹配")
    # 转为 JSON 友好的 datetime 字符串
    if job.get("finished_at"):
        job["elapsed_s"] = round(job["finished_at"] - (job.get("started_at") or job["created_at"]), 1)
    return job


@kbd_router.post("/{kbd_id}/reanalyze-image/{seq}", summary="重新识图单张图片（默认异步提交）")
async def reanalyze_single_image(request: Request, kbd_id: int, seq: int):
    """从 kbd_image 表读取指定 seq 的原始图片，重新识图。

    场景：用户在 admin-ui 图片列表中点击单张图片的刷新按钮，
    仅重新识图该图片，不影响其他图片。

    P1-8 修复：默认异步提交（202 + job_id），与批量 reanalyze-images 一致，
    避免单图 ~60s 的 LLM 调用长时间阻塞 HTTP 请求（网关超时风险）。
    向后兼容：?sync=true 走同步旧路径，直接返回完整结果（供 debug）。

    Returns（异步）:
        {"success": True, "kbd_id": int, "seq": int, "job_id": str, "status": "pending",
         "message": "Vision 单图任务已提交，请通过 GET /reanalyze-images/status 轮询进度"}
    """
    _check_auth(request)

    if _db_manager is None:
        raise HTTPException(status_code=503, detail="数据库未就绪")

    async with _db_manager.async_session_factory() as session:
        await _require_directly_mutable_kbd(session, kbd_id)

    trace_id = get_current_trace_id()
    logger.info(
        event="kbd_reanalyze_single_image_request",
        kbd_id=kbd_id,
        seq=seq,
        trace_id=trace_id,
    )

    # 向后兼容：?sync=true 直接同步返回（供 debug 场景）
    sync_mode = request.query_params.get("sync", "").lower() == "true"
    if sync_mode:
        return await _reanalyze_single_image_sync(kbd_id, seq, trace_id)

    # 异步：提交 Job（复用批量任务的状态机与轮询端点）
    from app.services.vision_job_manager import get_job_manager
    from app.services.vision_processor import reanalyze_single_image as do_reanalyze_single

    async def _runner(k_id: int) -> dict[str, Any]:
        # k_id 仅用于匹配 submit 签名；实际只处理当前 seq；透传 trace_id 串联日志
        return await do_reanalyze_single(k_id, seq, _db_manager.async_session_factory, trace_id=trace_id)

    jm = get_job_manager()
    job_id = await jm.submit(kbd_id, _runner, trace_id=trace_id)

    logger.info(
        event="kbd_reanalyze_single_image_submitted",
        kbd_id=kbd_id,
        seq=seq,
        job_id=job_id,
        trace_id=trace_id,
    )

    return JSONResponse(
        status_code=status.HTTP_202_ACCEPTED,
        content={
            "success": True,
            "kbd_id": kbd_id,
            "seq": seq,
            "job_id": job_id,
            "status": "pending",
            "message": "Vision 单图任务已提交，请通过 GET /reanalyze-images/status 轮询进度",
        },
    )


async def _reanalyze_single_image_sync(kbd_id: int, seq: int, trace_id: str):
    """旧同步路径（?sync=true 走此分支），保留兼容性。"""
    from app.services.vision_processor import reanalyze_single_image as do_reanalyze_single

    try:
        result = await do_reanalyze_single(kbd_id, seq, _db_manager.async_session_factory)
    except PublishedKbdMutationError as exc:
        raise HTTPException(
            status_code=409,
            detail={"code": "KBD_MAINTENANCE_WORKING_REQUIRED", "message": str(exc)},
        ) from exc
    except ValueError as exc:
        raise HTTPException(status_code=404, detail=str(exc))
    except RuntimeError as exc:
        raise HTTPException(status_code=503, detail=str(exc))
    except Exception as exc:
        logger.error(
            event="kbd_reanalyze_single_image_failed",
            kbd_id=kbd_id,
            seq=seq,
            error=str(exc),
            trace_id=trace_id,
        )
        raise HTTPException(status_code=500, detail=f"识图失败：{exc}")

    logger.info(
        event="kbd_reanalyze_single_image_completed",
        kbd_id=kbd_id,
        seq=seq,
        screenshot_type=result["screenshot_type"],
        trace_id=trace_id,
    )

    return {
        "success": True,
        "kbd_id": kbd_id,
        "seq": seq,
        "screenshot_type": result["screenshot_type"],
        "background": result["background"],
        "full_text": result["full_text"],
        "description": result["description"],
        "desc": result["desc"],
        "message": "识图完成",
    }


# ─────────────────────────────────────────────────────────────────────────────
# SOP 文档上传（docx 文件直接导入）
# ─────────────────────────────────────────────────────────────────────────────


def _collect_leaves(node: SOPNode) -> list[SOPNode]:
    """递归收集决策树的所有叶节点（children 为空的节点）。

    Args:
        node: SOPNode 根节点或子节点

    Returns:
        叶节点列表（按遍历顺序）
    """
    leaves: list[SOPNode] = []
    if not node.children:
        # 叶节点：无子节点
        leaves.append(node)
    else:
        # 中间节点：递归遍历子节点
        for child in node.children:
            leaves.extend(_collect_leaves(child))
    return leaves


def _parse_docx_bytes(content: bytes) -> tuple[str, str, list[tuple[str, str]]]:
    """解析 .docx 二进制内容，返回 (title, full_markdown, chapters)"""
    try:
        from docx import Document  # noqa: PLC0415
    except ImportError as exc:
        raise HTTPException(status_code=503, detail="服务器未安装 python-docx，请联系管理员") from exc

    doc = Document(io.BytesIO(content))

    title = ""
    md_lines: list[str] = []
    chapters: list[tuple[str, str]] = []
    current_chapter_title = "概述"
    current_chapter_lines: list[str] = []

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        text = para.text.strip()
        if not text:
            continue

        if style_name.startswith("Heading"):
            try:
                level = int(style_name.split()[-1])
            except ValueError:
                level = 1

            if current_chapter_lines:
                chapter_content = "\n".join(current_chapter_lines).strip()
                if chapter_content:
                    chapters.append((current_chapter_title, chapter_content))

            # 为保持与下游 _split_md_chapters 仅按 H1-H3 分章的既有契约一致，
            # 这里将 docx 的深层标题映射为最多三级 Markdown 标题，避免 H4+ 被写入
            # content_md 后无法被后续分块逻辑识别。
            heading_level = min(level, 3)
            heading_prefix = "#" * heading_level
            heading_line = f"{heading_prefix} {text}"
            md_lines.append(heading_line)
            current_chapter_title = text
            current_chapter_lines = [heading_line]

            if level == 1 and not title:
                title = text
        else:
            md_lines.append(text)
            current_chapter_lines.append(text)

    if current_chapter_lines:
        chapter_content = "\n".join(current_chapter_lines).strip()
        if chapter_content:
            chapters.append((current_chapter_title, chapter_content))

    if not title:
        title = "未命名 SOP 文档"

    full_markdown = "\n\n".join(md_lines)
    return title, full_markdown, chapters


def _split_md_chapters(content_md: str) -> list[tuple[str, str]]:
    """按 Markdown 标题分块，并合并无正文的标题章节到后续章节"""
    heading_pattern = re.compile(r"^(#{1,3})\s+(.+)$")
    current_title = "概述"
    current_lines: list[str] = []
    raw_chapters: list[tuple[str, str]] = []

    for line in content_md.split("\n"):
        match = heading_pattern.match(line)
        if match:
            if current_lines:
                content = "\n".join(current_lines).strip()
                if content:
                    raw_chapters.append((current_title, content))
            current_title = match.group(2).strip()
            current_lines = [line]
        else:
            current_lines.append(line)

    if current_lines:
        content = "\n".join(current_lines).strip()
        if content:
            raw_chapters.append((current_title, content))

    # 后处理：将无正文内容（仅含标题行）的章节合并到下一有正文章节
    def _has_body(text: str) -> bool:
        return any(line.strip() and not line.strip().startswith("#") for line in text.split("\n"))

    merged: list[tuple[str, str]] = []
    pending_content = ""
    pending_title = ""

    for title, content in raw_chapters:
        if _has_body(content):
            if pending_content:
                # 将无正文前缀并入当前有正文章节
                merged.append((pending_title, (pending_content + "\n\n" + content).strip()))
                pending_content = ""
                pending_title = ""
            else:
                merged.append((title, content))
        else:
            # 无正文章节，积累为后续章节前缀
            pending_content = (pending_content + "\n\n" + content).strip() if pending_content else content
            pending_title = title

    # 末尾残留的无正文章节（孤立标题）保留
    if pending_content:
        merged.append((pending_title, pending_content))

    return merged if merged else raw_chapters


@sop_router.post("/upload", status_code=status.HTTP_201_CREATED)
async def upload_sop_document(
    request: Request,
    file: UploadFile = File(..., description=".docx 或 .md 文件"),
    category_id: str | None = Form(None, description="分类编码，如 虚拟机-003"),
):
    """直接上传 .docx 或 .md 文件，解析后写入 SOP 草稿

    支持幂等：相同文件内容（SHA256 哈希）不会重复导入。
    上传成功后状态为 draft，需在本页面点击「发布」后方可被 AI 搜索。
    """
    _check_auth(request)

    if _db_manager is None:
        raise HTTPException(status_code=503, detail="数据库未就绪")

    filename = file.filename or ""
    file_ext = filename.lower().split(".")[-1] if "." in filename else ""

    if file_ext not in ("docx", "md"):
        raise HTTPException(status_code=400, detail="仅支持 .docx 或 .md 格式文件")

    content = await file.read()
    if len(content) > 50 * 1024 * 1024:  # 50MB 限制
        raise HTTPException(status_code=400, detail="文件过大，最大支持 50MB")

    file_hash = hashlib.sha256(content).hexdigest()

    # 根据文件类型解析
    try:
        if file_ext == "docx":
            doc_title, content_md, _ = _parse_docx_bytes(content)
        else:  # .md 文件
            content_md = content.decode("utf-8")
            # 从文件名或首行提取标题
            doc_title = filename.rsplit(".", 1)[0] if filename else "未命名 SOP"
            first_line = content_md.split("\n", 1)[0].strip()
            if first_line.startswith("# "):
                doc_title = first_line[2:].strip()
    except Exception as exc:
        logger.error(event="sop_upload_parse_error", filename=filename, error=str(exc))
        raise HTTPException(status_code=400, detail=f"文件解析失败：{exc}") from exc

    docx_hash = file_hash  # .docx 和 .md 均生成内容哈希，支持幂等导入

    async with _db_manager.async_session_factory() as session:
        # 幂等：已存在相同哈希则返回已有文档
        from sqlalchemy import select as sa_select  # noqa: PLC0415

        existing = await session.execute(sa_select(SopDocument).where(SopDocument.docx_hash == docx_hash))
        existing_doc = existing.scalar_one_or_none()
        if existing_doc:
            return {
                "success": True,
                "document_id": existing_doc.id,
                "status": existing_doc.status,
                "duplicate": True,
                "message": f"文件已导入（document_id={existing_doc.id}），跳过重复入库",
            }

        # 新建 sop_document
        sop_doc = SopDocument(
            source_id=f"sop-upload-{file_hash[:12]}",
            title=doc_title,
            content_md=content_md,
            category_id=category_id or None,
            docx_hash=docx_hash,
            status="draft",
        )
        session.add(sop_doc)
        await session.flush()

        await session.commit()

        document_id = sop_doc.id

    logger.info(
        event="sop_upload_completed",
        document_id=document_id,
        title=doc_title[:50],
        filename=filename,
    )
    return {
        "success": True,
        "document_id": document_id,
        "status": "draft",
        "duplicate": False,
        "title": doc_title,
    }
